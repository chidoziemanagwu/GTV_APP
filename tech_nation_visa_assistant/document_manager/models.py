from django.db import models
from accounts.models import User
from django.db.models.signals import post_save
from django.dispatch import receiver
from django.utils import timezone
import logging

logger = logging.getLogger(__name__)




class EligibilityCriteria(models.Model):
    TYPE_CHOICES = (
        ('mandatory', 'Mandatory'),
        ('optional', 'Optional'),
    )
    PATH_CHOICES = (
        ('talent', 'Exceptional Talent'),
        ('promise', 'Exceptional Promise'),
        ('both', 'Both Paths'),
    )

    name = models.CharField(max_length=200)
    description = models.TextField()
    criteria_type = models.CharField(max_length=10, choices=TYPE_CHOICES)
    applicable_path = models.CharField(max_length=10, choices=PATH_CHOICES)
    number_of_documents = models.IntegerField(default=2)
    notion_link = models.URLField(blank=True)

    def __str__(self):
        return f"{self.criteria_type.title()} - {self.name}"

# document_manager/models.py

class Document(models.Model):
    STATUS_CHOICES = (
        ('not_started', 'Not Started'),
        ('draft', 'Draft'),
        ('in_progress', 'In Progress'),
        ('reviewing', 'Under Review'),
        ('completed', 'Completed'),
        ('archived', 'Archived'),
    )

    TYPE_CHOICES = (
        ('personal_statement', 'Personal Statement'),
        ('cv', 'CV'),
        ('recommendation', 'Recommendation Letter'),
        ('evidence', 'Evidence Document'),
        ('other', 'Other'),
    )

    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='documents')
    title = models.CharField(max_length=255)
    document_type = models.CharField(max_length=50, choices=TYPE_CHOICES)
    related_criteria = models.ForeignKey(
        EligibilityCriteria,
        on_delete=models.SET_NULL,
        null=True,
        blank=True
    )

    # Content and File
    content = models.TextField(blank=True, null=True)
    file = models.FileField(upload_to='user_documents/', null=True, blank=True)

    # Status and Metadata
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='not_started')
    is_generated = models.BooleanField(default=False)
    generation_prompt = models.TextField(blank=True, null=True)

    # Document Analysis
    word_count = models.IntegerField(default=0)
    page_count = models.IntegerField(default=0)

    # Timestamps
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    source_cv = models.ForeignKey(
        'self',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='generated_documents',
        help_text="If this is a personal statement, links to the CV used to generate it"
    )

    original_file = models.FileField(
        upload_to='original_documents/',
        null=True,
        blank=True,
        help_text="Original uploaded document"
    )

    generated_file = models.FileField(
        upload_to='generated_documents/',
        null=True,
        blank=True,
        help_text="Generated document in DOCX format"
    )

    is_chosen = models.BooleanField(
        default=False,
        help_text="Indicates if this document is the user's chosen version for submission"
    )

    notes = models.TextField(blank=True, null=True)  # Add this line
    
    class Meta:
        ordering = ['-updated_at']

    def __str__(self):
        return f"{self.title} - {self.get_document_type_display()}"

    @property
    def has_original_file(self):
        return bool(self.original_file)

    def generate_docx(self):
        """Generate a DOCX file with disclaimer"""
        from docx import Document as DocxDocument

        doc = DocxDocument()

        # Add disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer.add_run("DISCLAIMER").bold = True
        doc.add_paragraph(
            "This document is auto-generated for sample purposes only. "
            "It should be thoroughly reviewed, customized, and edited to reflect "
            "your personal experiences and achievements. Do not submit this "
            "document as-is for your visa application."
        )

        # Add a divider
        doc.add_paragraph("=" * 50)

        # Add the main content
        doc.add_paragraph(self.content)

        # Save the file
        filename = f"{self.title}_{self.id}.docx"
        filepath = f"generated_documents/{filename}"
        doc.save(filepath)

        # Update the generated_file field
        self.generated_file.name = filepath
        self.save()

    def get_file_extension(self):
        if self.file:
            return self.file.name.split('.')[-1].lower()
        return None



class DocumentComment(models.Model):
    document = models.ForeignKey(Document, on_delete=models.CASCADE, related_name='comments')
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    content = models.TextField()
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"Comment on {self.document.title}"
    



class UserPoints(models.Model):
    user = models.OneToOneField(User, on_delete=models.CASCADE, related_name='points')
    balance = models.IntegerField(default=0)
    lifetime_points = models.IntegerField(default=0) # Total points ever acquired
    last_purchase = models.DateTimeField(null=True, blank=True)

    def __str__(self):
        return f"{self.user.username} - {self.balance} points"

    def add_points(self, amount):
        if amount <= 0: # Prevent adding zero or negative points if that's not desired
            logger.warning(f"Attempted to add non-positive points ({amount}) for user {self.user.username}")
            return

        self.balance += amount
        self.lifetime_points += amount
        self.save() # Save UserPoints first

        # Update the user's profile's is_paid_user status
        try:
            profile = self.user.profile # Assuming User has a 'profile' one-to-one link to UserProfile
            if not profile.is_paid_user and self.balance > 0:
                profile.is_paid_user = True
                profile.save(update_fields=['is_paid_user'])
        except User.profile.RelatedObjectDoesNotExist: # More specific exception
            logger.error(f"UserProfile not found for user {self.user.username} during add_points.")
        except Exception as e:
            logger.error(f"Error updating profile's is_paid_user status for {self.user.username} during add_points: {e}")


    def use_points(self, amount, description=None): # MODIFIED LINE
        if amount <= 0: # Prevent using zero or negative points
            logger.warning(f"Attempted to use non-positive points ({amount}) for user {self.user.username}. Description: {description}")
            return False
            
        if self.balance >= amount:
            self.balance -= amount
            self.save(update_fields=['balance']) # Explicitly save only balance here for clarity

            # Log the point usage with description
            logger.info(f"Successfully used {amount} points for user {self.user.username}. Purpose: {description}. New balance: {self.balance}")

            # Update the user's profile's is_paid_user status if necessary
            try:
                profile = self.user.profile
                if profile.is_paid_user and self.balance <= 0:
                    profile.is_paid_user = False
                    profile.save(update_fields=['is_paid_user'])
            except User.profile.RelatedObjectDoesNotExist:
                logger.error(f"UserProfile not found for user {self.user.username} during use_points.")
            except Exception as e:
                logger.error(f"Error updating profile's is_paid_user status for {self.user.username} during use_points: {e}")
            
            # Optional: Create a PointTransaction record here if you want a detailed history
            # For example:
            # PointTransaction.objects.create(
            #     user=self.user,
            #     points=-amount, # Store as negative for usage
            #     description=description,
            #     # ... other relevant fields for a usage transaction
            # )

            return True
        else:
            logger.warning(f"Failed to use {amount} points for user {self.user.username}. Insufficient balance. Current balance: {self.balance}. Attempted for: {description}")
            return False
        

        
class PointsPackage(models.Model):
    name = models.CharField(max_length=100)
    points = models.IntegerField()
    price = models.DecimalField(max_digits=6, decimal_places=2)
    description = models.TextField()
    is_active = models.BooleanField(default=True)

    def __str__(self):
        return f"{self.name} - {self.points} points (£{self.price})"
    

class PointsTransaction(models.Model):
    PAYMENT_STATUS_CHOICES = (
        ('pending', 'Pending'),
        ('completed', 'Completed'),
        ('failed', 'Failed'),
        ('refunded', 'Refunded'),
    )

    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='points_transactions')
    package = models.ForeignKey(PointsPackage, on_delete=models.SET_NULL, null=True)
    amount = models.DecimalField(max_digits=10, decimal_places=2)
    points = models.IntegerField()
    payment_status = models.CharField(max_length=20, choices=PAYMENT_STATUS_CHOICES, default='pending')
    stripe_payment_intent_id = models.CharField(max_length=255, blank=True, null=True)
    stripe_checkout_id = models.CharField(max_length=255, blank=True, null=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    def __str__(self):
        return f"{self.user.username} - {self.points} points - {self.get_payment_status_display()}"


# --- SIGNALS ---

@receiver(post_save, sender=UserPoints)
def sync_is_paid_user_on_profile(sender, instance, **kwargs):
    """
    Syncs the is_paid_user status on UserProfile based on UserPoints balance.
    This is somewhat redundant if add_points and use_points handle it,
    but can act as a failsafe or handle direct balance manipulations if any.
    """
    try:
        profile = instance.user.profile # instance is UserPoints
        new_is_paid_status = instance.balance > 0
        if profile.is_paid_user != new_is_paid_status:
            profile.is_paid_user = new_is_paid_status
            profile.save(update_fields=['is_paid_user'])
            logger.info(f"Synced is_paid_user for {instance.user.username} to {new_is_paid_status} via UserPoints signal.")
    except User.profile.RelatedObjectDoesNotExist:
        logger.error(f"UserProfile not found for user {instance.user.username} during UserPoints signal sync.")
    except Exception as e:
        logger.error(f"Error in sync_is_paid_user_on_profile for {instance.user.username}: {e}")


@receiver(post_save, sender='accounts.UserProfile') # Sender is a string 'app_label.ModelName'
def ensure_user_points_exists_on_profile_creation(sender, instance, created, **kwargs):
    """
    Ensures a UserPoints record exists when a UserProfile is created.
    The initial point balance (e.g., 3 for new users) should be handled by
    the signal connected to User creation in accounts/models.py.
    This signal primarily ensures the UserPoints object is present.
    """
    if created: # instance here is a UserProfile
        # UserPoints model is defined in this file, so direct reference is fine.
        user_points, up_created = UserPoints.objects.get_or_create(user=instance.user)
        if up_created:
            # This means the signal in accounts/models.py (listening to User creation)
            # might not have run yet or didn't create UserPoints for some reason.
            # UserPoints.balance defaults to 0 as per model definition.
            logger.info(f"UserPoints record created for {instance.user.username} via UserProfile creation signal. Balance defaults to {user_points.balance}.")
        # else:
            # logger.info(f"UserPoints record already existed for {instance.user.username} when UserProfile was created.")


# --- REFERRAL POINTS LOGIC (Seems okay, but let's review its interaction) ---

def award_referral_points(user_who_made_purchase): # Renamed for clarity
    """
    Awards referral rewards (now 1 free use, not points) to the referrer
    when a user makes their first qualifying purchase/action.
    This function is called externally, e.g., after a successful payment.
    """
    from referrals.models import ReferralSignup # Correctly imported locally

    try:
        # Find if this user was referred by someone and the reward hasn't been processed
        referral_signup = ReferralSignup.objects.get(
            referred_user=user_who_made_purchase,
            has_been_rewarded=False # Ensure we only process once
        )

        # The award_rewards method on ReferralSignup now handles granting the free use
        # and marking itself as rewarded.
        if referral_signup.award_rewards(): # This method now returns True on success
            logger.info(f"Successfully processed referral reward (1 free use to referrer) for referred user {user_who_made_purchase.username} via award_referral_points function.")
            return True
        else:
            logger.warning(f"ReferralSignup.award_rewards() failed or indicated no action for referred user {user_who_made_purchase.username}.")
            return False

    except ReferralSignup.DoesNotExist:
        logger.info(f"User {user_who_made_purchase.username} made a purchase but was not found in ReferralSignup or reward already processed.")
        return False
    except Exception as e:
        logger.error(f"Error in award_referral_points for {user_who_made_purchase.username}: {e}", exc_info=True)
        return False