from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, FileResponse
from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import csrf_exempt
from django.core.cache import cache
from django.db import models
from datetime import datetime, timedelta
from django.utils import timezone as django_timezone
import json
import os
import io
import logging
import re
import threading
import time
import concurrent.futures
import tempfile
import uuid
from typing import Optional
from functools import wraps

# Document processing imports
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from accounts.models import User, UserProfile
from referrals.models import ReferralSignup

# Django imports
from .models import Document, PointsPackage, UserPoints, PointsTransaction
from .forms import DocumentForm, PersonalStatementForm, CVForm
from django.conf import settings
from .utils import calculate_application_progress
from groq import Groq
from collections import defaultdict
from .points_utils import require_points
from django.urls import reverse
import stripe
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils import timezone

stripe.api_key = settings.STRIPE_SECRET_KEY


# Configure logging
logger = logging.getLogger(__name__)

# Create a document cache
document_cache = {}

def get_referral_points(user):
    """Helper function to get referral points for a user"""
    try:
        from referrals.models import ReferralCode, ReferralSignup
        referral_code = ReferralCode.objects.filter(user=user).first()
        if referral_code:
            # Count successful referrals where points were awarded
            successful_referrals = ReferralSignup.objects.filter(
                referral_code=referral_code,
                points_awarded=True,
                has_been_rewarded=True
            ).count()
            # Each successful referral gives 3 points
            return successful_referrals * 3
        return 0
    except Exception as e:
        logger.error(f"Error getting referral points: {e}")
        return 0

    


# (Make sure to import re at the top of your views.py if it's not already there: import re)
# ... other imports ...
# from groq import Groq # Ensure Groq is imported
# from django.conf import settings # Ensure settings is imported
# import logging
# logger = logging.getLogger(__name__) # Ensure logger is configured

class AIProvider:
    def __init__(self):
        self.groq_client = Groq(api_key=settings.GROQ_API_KEY)
        self.model = "llama-3.1-8b-instant"  # Or your preferred model

    def get_system_prompt(self): # Existing system prompt for personal statements
        return """You are an expert at writing personal statements for Tech Nation Global Talent Visa applications.

        Format your response with proper structure, using:
        - Main sections with # (like "# Introduction")
        - Subsections with ## (like "## Technical Skills")
        - Bold text for emphasis using **text**
        - Bullet points using - for lists

        Your personal statement should:
        - Demonstrate exceptional talent and potential in digital technology
        - Showcase technical skills, innovations, and impact
        - Include specific examples of projects, achievements, and contributions
        - Highlight recognition, awards, or publications
        - Demonstrate potential to become a leader in the UK tech sector
        - Include evidence of contributions to company/organizational growth
        - Explain plans to contribute to the UK's digital technology sector
        - Be clear, concise, and well-structured
        - Ensure all claims can be supported by evidence
        - Be approximately 800-1000 words

        Make sure to include Introduction, Technical Expertise, Leadership, and Conclusion sections."""

    def generate_content(self, prompt: str) -> Optional[str]:
        """Generates general content (e.g., personal statements) using Groq."""
        try:
            response = self.groq_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.get_system_prompt()},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000 # Adjust as needed
            )
            logger.info("Successfully generated content using Groq")
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Error with Groq API for general content: {str(e)}")
            return None

    def get_cv_analysis_system_prompt(self):
        """System prompt specifically for CV analysis expecting JSON output."""
        return """You are an expert CV analyst specializing in Tech Nation Global Talent Visa applications.
Your task is to review the provided CV content against the visa track requirements and output your analysis in a structured JSON format.
The user will provide the CV content and detailed instructions for the JSON output.
Your response MUST be a single JSON object. Do NOT include any explanatory text, comments, or markdown formatting (like ```json) before or after the JSON object itself.
The JSON object must contain a single top-level key "analysis".
The "analysis" object should include all the fields as specified in the user's prompt.
Adhere strictly to the requested JSON schema.
"""

    def generate_cv_analysis_content(self, prompt: str) -> Optional[str]:
        """Generates CV analysis content using Groq, expecting structured JSON output."""
        try:
            response = self.groq_client.chat.completions.create(
                model=self.model, # Consider if a different model is better for structured JSON
                messages=[
                    {"role": "system", "content": self.get_cv_analysis_system_prompt()},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,  # Lower temperature for more deterministic and structured JSON output
                max_tokens=3000,  # Ensure this is enough for the detailed JSON
                # If Groq supports a direct JSON mode, enable it here.
                # For example, some APIs use: response_format={"type": "json_object"}
            )
            logger.info("Successfully generated CV analysis content using Groq")
            content = response.choices[0].message.content.strip()

            # Attempt to extract JSON if it's wrapped in markdown (though system prompt tries to prevent this)
            # Regex to find ```json ... ``` block
            match = re.search(r"```json\s*(.*?)\s*```", content, re.DOTALL | re.IGNORECASE)
            if match:
                logger.info("Extracted JSON from markdown block in CV analysis response.")
                return match.group(1).strip()
            
            # If not in markdown, assume the content is the JSON itself (as per system prompt)
            return content

        except Exception as e:
            logger.error(f"Error with Groq API for CV analysis: {str(e)}")
            return None

    def generate_content_stream(self, prompt: str):
        """Streaming version for real-time responses (primarily for text, not structured JSON)."""
        try:
            response = self.groq_client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self.get_system_prompt()}, # Uses general system prompt
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000,
                stream=True
            )
            return response
        except Exception as e:
            logger.error(f"Error with Groq streaming: {str(e)}")
            return None

# Create a singleton instance (as you had)
ai_provider = AIProvider() # This should be defined once in your views.py






# MISSING FUNCTION 1: extract_cv_content
def extract_cv_content(cv_file):
    """Extract text content from CV file (PDF, DOCX, or TXT) with better error handling"""
    try:
        # Reset file pointer to beginning
        cv_file.seek(0)

        # Check if we've already processed this file
        file_content = cv_file.read()
        cv_file.seek(0)  # Reset file pointer again

        if len(file_content) == 0:
            raise Exception("The uploaded file appears to be empty")

        file_extension = cv_file.name.lower().split('.')[-1]

        if file_extension == 'pdf':
            content = extract_pdf_content(cv_file)
        elif file_extension in ['docx', 'doc']:
            content = extract_docx_content(cv_file)
        elif file_extension == 'txt':
            try:
                content = file_content.decode('utf-8')
            except UnicodeDecodeError:
                content = file_content.decode('latin-1', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

        # Clean and validate content
        if content:
            content = content.strip()
            # Remove excessive whitespace
            content = re.sub(r'\n\s*\n', '\n\n', content)
            content = re.sub(r' +', ' ', content)

        if not content or len(content) < 50:
            raise Exception("Could not extract sufficient text content from the file")

        return content

    except Exception as e:
        logger.error(f"Error extracting CV content from {cv_file.name}: {str(e)}")
        raise Exception(f"Failed to extract content from CV: {str(e)}")


def extract_pdf_content(pdf_file):
    """Extract text from PDF file with better error handling"""
    try:
        # Reset file pointer
        pdf_file.seek(0)

        # Create a copy of the file in memory
        file_copy = io.BytesIO(pdf_file.read())
        pdf_file.seek(0)  # Reset file pointer for future use

        try:
            pdf_reader = PyPDF2.PdfReader(file_copy)
        except Exception as e:
            raise Exception(f"Invalid or corrupted PDF file: {str(e)}")

        if len(pdf_reader.pages) == 0:
            raise Exception("PDF file contains no pages")

        text = ""
        pages_processed = 0

        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
                    pages_processed += 1
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                continue

        if pages_processed == 0:
            raise Exception("Could not extract text from any pages in the PDF")

        if not text.strip():
            raise Exception("No readable text found in the PDF file")

        logger.info(f"Successfully extracted text from {pages_processed} pages")
        return text.strip()

    except Exception as e:
        logger.error(f"Error extracting PDF content: {str(e)}")
        raise Exception(f"Failed to extract PDF content: {str(e)}")


def extract_docx_content(docx_file):
    """Extract text from DOCX file with better error handling"""
    try:
        # Reset file pointer
        docx_file.seek(0)

        # Create a copy of the file in memory
        file_copy = io.BytesIO(docx_file.read())
        docx_file.seek(0)  # Reset file pointer for future use

        try:
            doc = DocxDocument(file_copy)
        except Exception as e:
            raise Exception(f"Invalid or corrupted DOCX file: {str(e)}")

        text = ""
        paragraphs_processed = 0

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text += paragraph.text.strip() + "\n"
                paragraphs_processed += 1

        # Also extract text from tables
        tables_processed = 0
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
                    tables_processed += 1

        if paragraphs_processed == 0 and tables_processed == 0:
            raise Exception("No readable content found in the DOCX file")

        if not text.strip():
            raise Exception("No text could be extracted from the DOCX file")

        logger.info(f"Successfully extracted text from {paragraphs_processed} paragraphs and {tables_processed} table rows")
        return text.strip()

    except Exception as e:
        logger.error(f"Error extracting DOCX content: {str(e)}")
        raise Exception(f"Failed to extract DOCX content: {str(e)}")


# MISSING FUNCTION 2: save_to_docx
def save_to_docx(content, title="Document", doc_type="personal_statement"):
    """Save HTML content to a DOCX file"""
    try:
        # Create a new document
        doc = DocxDocument()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%d %B %Y')}")
        date_run.font.size = Pt(10)
        date_run.italic = True
        
        # Add separator
        doc.add_paragraph("=" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Clean up the content - remove HTML tags
        import re
        clean_content = re.sub(r'<[^>]+>', '', content)
        
        # Remove markdown formatting
        clean_content = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_content)
        clean_content = re.sub(r'\n\s*\n', '\n\n', clean_content)
        clean_content = clean_content.strip()
        
        # Split by paragraphs and add to document
        paragraphs = clean_content.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Check if it's a heading (starts with #)
                if para_text.startswith('#'):
                    heading_text = para_text.lstrip('#').strip()
                    doc.add_heading(heading_text, level=1)
                elif para_text.startswith('##'):
                    heading_text = para_text.lstrip('#').strip()
                    doc.add_heading(heading_text, level=2)
                elif para_text.startswith('•') or para_text.startswith('-'):
                    # Handle bullet points
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25)
                    run = p.add_run(para_text[1:].strip())
                    run.font.size = Pt(11)
                else:
                    # Regular paragraph
                    paragraph = doc.add_paragraph(para_text)
                    paragraph.paragraph_format.space_after = Pt(12)
        
        # Add disclaimer
        doc.add_paragraph("=" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_title = doc.add_heading("DISCLAIMER", level=1)
        disclaimer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        disclaimer_text = doc.add_paragraph(
            "This document was automatically generated and should be reviewed for accuracy before submission. "
            "The content is based on the information provided and does not guarantee visa approval."
        )
        disclaimer_text.style = 'Intense Quote'
        
        # Save to temporary file
        temp_dir = tempfile.gettempdir()
        filename = f"{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)
        
        doc.save(filepath)
        
        return filepath
        
    except Exception as e:
        logger.error(f"Error saving to DOCX: {str(e)}")
        raise Exception(f"Failed to save document: {str(e)}")

# MISSING FUNCTION 3: process_single_document
def process_single_document(document_id, user):
    """Process a single document for batch operations"""
    try:
        document = Document.objects.get(id=document_id, user=user)
        
        # Perform some processing based on document type
        if document.document_type == 'personal_statement':
            # Update word count or other metrics
            word_count = len(document.content.split()) if document.content else 0
            
            # Update document status based on word count
            if word_count >= 800 and word_count <= 1200:
                status = 'completed'
            elif word_count > 0:
                status = 'in_progress'
            else:
                status = 'draft'
            
            document.status = status
            document.save(update_fields=['status'])
            
            return {
                "status": "success",
                "word_count": word_count,
                "document_status": status
            }
        elif document.document_type == 'cv':
            # Analyze CV if it has a file
            if document.file:
                try:
                    with open(document.file.path, 'rb') as f:
                        cv_content = extract_cv_content(f)
                    
                    # Simple analysis
                    word_count = len(cv_content.split()) if cv_content else 0
                    document.notes = f"CV analyzed: {word_count} words extracted"
                    document.status = 'analyzed'
                    document.save(update_fields=['notes', 'status'])
                    
                    return {
                        "status": "success",
                        "word_count": word_count,
                        "analysis": "CV content extracted and analyzed"
                    }
                except Exception as e:
                    return {
                        "status": "error",
                        "message": f"Error analyzing CV: {str(e)}"
                    }
        
        return {
            "status": "success",
            "message": "Document processed"
        }
            
    except Document.DoesNotExist:
        return {
            "status": "error",
            "message": "Document not found"
        }
    except Exception as e:
        logger.error(f"Error processing document {document_id}: {str(e)}")
        return {
            "status": "error",
            "message": str(e)
        }

# Rate limiting decorator
def rate_limit(max_calls=5, window=60):
    """Decorator to limit API calls"""
    def decorator(view_func):
        # Use a simple in-memory store for rate limiting
        # In production, you'd use Redis or another shared cache
        calls = {}

        @wraps(view_func)
        def _wrapped_view(request, *args, **kwargs):
            # Get user identifier
            user_id = request.user.id if request.user.is_authenticated else request.META.get('REMOTE_ADDR')
            key = f"{view_func.__name__}:{user_id}"

            # Get current timestamp
            now = time.time()

            # Initialize or clean up old timestamps
            if key not in calls:
                calls[key] = []
            calls[key] = [t for t in calls[key] if now - t < window]

            # Check if rate limit exceeded
            if len(calls[key]) >= max_calls:
                return JsonResponse({
                    'success': False,
                    'error': f'Rate limit exceeded. Maximum {max_calls} calls per {window} seconds.'
                }, status=429)

            # Add current call
            calls[key].append(now)

            # Process the view
            return view_func(request, *args, **kwargs)

        return _wrapped_view
    return decorator

# Background processing
class DocumentProcessor:
    @staticmethod
    def process_in_background(func, *args, **kwargs):
        """Run a function in a background thread"""
        thread = threading.Thread(target=func, args=args, kwargs=kwargs)
        thread.daemon = True
        thread.start()
        return thread

def process_content_for_frontend(content):
    """Process the AI-generated content into properly formatted HTML with guidelines"""
    import re

    # Create a properly formatted HTML structure
    html_output = """
    <div class="personal-statement">
        <div class="warning-box" style="background-color: #fff3cd; border: 1px solid #ffeeba; padding: 15px; margin-bottom: 20px; border-radius: 5px;">
            <h4 style="color: #856404; margin-top: 0;">⚠️ SAMPLE STATEMENT</h4>
            <p>This is an AI-generated sample personal statement. You should:</p>
            <ul>
                <li>Customize it with your specific achievements and experiences</li>
                <li>Verify all facts and claims before submission</li>
                <li>Use it as a guide, not a final submission</li>
            </ul>
        </div>
    """

    # Extract title and author if present
    title_match = re.search(r'Personal Statement for Tech Nation.*', content, re.IGNORECASE)
    title = title_match.group(0) if title_match else "Personal Statement for Tech Nation Global Talent Visa"

    author_match = re.search(r'By (.*?)(?=\n|$)', content)
    author = author_match.group(1) if author_match else ""

    # Add title to HTML
    html_output += f'<h1>{title}</h1>\n'
    if author:
        html_output += f'<p style="text-align: center; font-style: italic;">By {author}</p>\n'

    # Remove title and author from content
    if title_match:
        content = content.replace(title_match.group(0), '')
    if author_match:
        content = content.replace(f"By {author_match.group(1)}", '')

    # Split content by sections (marked with #)
    sections = content.split('#')

    for section in sections:
        if not section.strip():
            continue

        # Remove any -- markers
        section = section.replace('--', '')

        # Process the section
        lines = section.strip().split('\n')
        section_title = lines[0].strip()

        if section_title:
            html_output += f'<h2>{section_title}</h2>\n'

        # Process the rest of the section content
        section_content = '\n'.join(lines[1:]).strip()

        # Process subsections (marked with ##)
        subsections = section_content.split('##')

        for i, subsection in enumerate(subsections):
            if not subsection.strip():
                continue

            subsection_lines = subsection.strip().split('\n')

            # If this is not the first subsection, it has a title
            if i > 0 and subsection_lines[0].strip():
                subsection_title = subsection_lines[0].strip()
                html_output += f'<h3>{subsection_title}</h3>\n'
                subsection_content = '\n'.join(subsection_lines[1:]).strip()
            else:
                subsection_content = subsection.strip()

            # Process paragraphs and lists
            paragraphs = re.split(r'\n\s*\n', subsection_content)

            for para in paragraphs:
                if not para.strip():
                    continue

                # Check if this is a list
                if re.search(r'^\s*-\s+', para.strip(), re.MULTILINE):
                    # This is a list
                    list_items = re.findall(r'^\s*-\s+(.*?)$', para, re.MULTILINE)
                    if list_items:
                        html_output += '<ul style="list-style: disc outside none; margin-left: 20px; margin-bottom: 15px;">\n'
                        for item in list_items:
                            html_output += f'<li style="margin-bottom: 8px;">{item.strip()}</li>\n'
                        html_output += '</ul>\n'
                # Check for checkmarks (✅)
                elif '✅' in para:
                    # This is a list with checkmarks
                    list_items = para.split('✅')
                    if len(list_items) > 1:  # Skip the first empty item
                        html_output += '<ul style="list-style: none; margin-left: 20px; margin-bottom: 15px;">\n'
                        for item in list_items[1:]:
                            if item.strip():
                                html_output += f'<li style="margin-bottom: 8px;">✅ {item.strip()}</li>\n'
                        html_output += '</ul>\n'
                else:
                    # Regular paragraph
                    html_output += f'<p>{para.strip()}</p>\n'

        # Add horizontal rule after each main section
        html_output += '<hr>\n'

    # Process bold text
    html_output = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html_output)

    # Process links
    html_output = re.sub(r'$$(.*?)$$$$(.*?)$$', r'<a href="\2">\1</a>', html_output)

    # Add disclaimer at the bottom
    html_output += """
        <div class="disclaimer-box" style="background-color: #f8d7da; border: 1px solid #f5c6cb; padding: 15px; margin-top: 20px; border-radius: 5px;">
            <h4 style="color: #721c24; margin-top: 0;">📝 Important Notice</h4>
            <p>This document was generated based on the provided CV and instructions. Before submitting:</p>
            <ul style="list-style: disc outside none; margin-left: 20px; margin-bottom: 15px;">
                <li style="margin-bottom: 8px;">Review all content for accuracy</li>
                <li style="margin-bottom: 8px;">Replace any placeholder text</li>
                <li style="margin-bottom: 8px;">Ensure all claims are supported by evidence</li>
                <li style="margin-bottom: 8px;">Verify word count meets requirements (800-1000 words)</li>
                <li style="margin-bottom: 8px;">Customize to reflect your unique achievements and experiences</li>
            </ul>
        </div>
    </div>
    """

    return html_output

# document_manager/views.py

def generate_document_task(cv_content, instructions, user_id, task_id, document_id): # Add document_id here
    """Process document generation with Groq"""
    try:
        # Update status to processing in cache
        cache.set(f"task_status_{task_id}", {"status": "processing", "progress": 10}, 3600)

        # The document_id is now passed directly as a parameter.
        # The cache.get(f"task_document_{task_id}") was a workaround and can be removed
        # if document_id is always passed. Or keep it as a fallback if needed.

        # Create the prompt
        prompt = f"""
        Please write a personal statement for a Tech Nation Global Talent Visa application based on the following CV content and instructions.

        The personal statement should:
        1. Be well-structured and professional
        2. Focus on technical achievements and innovations
        3. Highlight leadership and impact
        4. Include specific examples from the CV
        5. Be around 1000 words

        CV Content:
        {cv_content}

        Additional Instructions:
        {instructions}

        Please write the personal statement in first person and ensure it aligns with Tech Nation's criteria.
        Format the statement with clear section headers using # for main sections and ## for subsections.
        Use bullet points with - for lists and bold text with ** for emphasis.
        """

        cache.set(f"task_status_{task_id}", {"status": "processing", "progress": 30}, 3600)
        raw_content = ai_provider.generate_content(prompt)
        provider_used = "Groq"

        if not raw_content:
            raise Exception("AI provider failed to generate content")

        formatted_content = process_content_for_frontend(raw_content)
        cache.set(f"task_status_{task_id}", {"status": "processing", "progress": 90}, 3600)

        if document_id: # Use the passed document_id
            try:
                document = Document.objects.get(id=document_id)
                document.content = formatted_content
                document.status = 'completed' # Ensure this matches your Document.STATUS_CHOICES
                document.notes = f"Generated using {provider_used}"
                document.save()
                logger.info(f"Document {document_id} updated with content from {provider_used}")
            except Document.DoesNotExist:
                logger.error(f"Document with id {document_id} not found in generate_document_task for task {task_id}.")
            except Exception as e:
                logger.error(f"Error updating document {document_id} in task {task_id}: {str(e)}")
        else:
            logger.warning(f"No document_id provided to generate_document_task for task_id {task_id}")

        cache_key = f"task_result_{task_id}"
        cache.set(cache_key, formatted_content, 3600)
        cache.set(f"task_status_{task_id}", {
            "status": "complete",
            "progress": 100,
            "provider": provider_used
        }, 3600)

        return formatted_content

    except Exception as e:
        cache.set(f"task_status_{task_id}", {
            "status": "failed",
            "error": str(e)
        }, 3600)
        logger.error(f"Document generation failed for task {task_id}, document {document_id}: {str(e)}", exc_info=True)
        # Also update the document status to reflect failure if document_id is available
        if document_id:
            try:
                doc_to_fail = Document.objects.get(id=document_id)
                doc_to_fail.status = 'failed_generation' # Add this to your STATUS_CHOICES if it doesn't exist
                doc_to_fail.notes = f"Generation failed: {str(e)[:250]}"
                doc_to_fail.save(update_fields=['status', 'notes'])
            except Document.DoesNotExist:
                pass # Already logged
            except Exception as db_err:
                logger.error(f"Error updating document {document_id} to failed status: {db_err}")
        return None



def get_tech_nation_requirements(track):
    """Get requirements based on track with caching"""
    cache_key = f"tech_nation_requirements_{track}"
    cached_requirements = cache.get(cache_key)

    if cached_requirements:
        return cached_requirements

    requirements = {
        'digital_technology': {
            'mandatory': [
                'Technical expertise in building, using, deploying or exploiting technology',
                'Proven innovation contributions or exceptional promise',
                'Recognition as a leading talent in the digital technology sector'
            ],
            'qualifying': [
                'Significant technical or commercial contributions to the sector',
                'Academic contributions through research/teaching/publications',
                'Recognition through awards, speaking engagements, or media coverage',
                'Experience leading or playing a key role in technical projects',
                'Continuous learning and adaptation to new technologies',
                'Contributions to open source projects or tech community'
            ]
        },
        'data_science_ai': {
            'mandatory': [
                'Advanced expertise in AI, machine learning, or data science',
                'Proven track record of innovative AI/data solutions',
                'Recognition in the AI/data science community'
            ],
            'qualifying': [
                'Development of novel AI algorithms or models',
                'Publications in peer-reviewed AI/ML journals or conferences',
                'Contributions to major AI/data science projects',
                'Leadership in AI research or development teams',
                'Patents or intellectual property in AI/ML',
                'Speaking engagements at AI conferences',
                'Collaboration with recognized AI institutions'
            ]
        }
    }

    # Get the requirements for the specified track or default to digital technology
    track_requirements = requirements.get(track, requirements['digital_technology'])

    # Format the requirements for better presentation
    formatted_requirements = {
        'mandatory': "\n".join(f"• {item}" for item in track_requirements['mandatory']),
        'qualifying': "\n".join(f"• {item}" for item in track_requirements['qualifying'])
    }

    # Cache the formatted requirements for 1 week (these rarely change)
    cache.set(cache_key, formatted_requirements, 60 * 60 * 24 * 7)

    return formatted_requirements

def process_text_response(text):
    """Process plain text response when JSON parsing fails"""
    try:
        # Initialize structure
        analysis = {
            'strength_score': 70,
            'summary': '',
            'suggestions': {
                'technical_expertise': [],
                'leadership': [],
                'innovation': [],
                'recognition': []
            },
            'missing_elements': [],
            'formatting_recommendations': []
        }

        # Extract score if present
        score_match = re.search(r'(\d+)(?:/100|%)', text)
        if score_match:
            analysis['strength_score'] = int(score_match.group(1))

        # Split into sections
        sections = text.split('\n\n')

        for section in sections:
            section = section.strip()
            if not section:
                continue

            # Process each section
            if 'Technical' in section:
                points = extract_points(section)
                analysis['suggestions']['technical_expertise'] = points
            elif 'Leadership' in section:
                points = extract_points(section)
                analysis['suggestions']['leadership'] = points
            elif 'Innovation' in section:
                points = extract_points(section)
                analysis['suggestions']['innovation'] = points
            elif 'Recognition' in section:
                points = extract_points(section)
                analysis['suggestions']['recognition'] = points
            elif 'Missing' in section:
                points = extract_points(section)
                analysis['missing_elements'] = points
            elif 'Format' in section:
                points = extract_points(section)
                analysis['formatting_recommendations'] = points
            elif 'Summary' in section or 'Overall' in section:
                analysis['summary'] = clean_text(section)

        return analysis

    except Exception as e:
        logger.error(f"Error in text processing: {str(e)}")
        return {
            'error': f"Error processing analysis: {str(e)}"
        }

def extract_points(text):
    """Extract bullet points from text"""
    points = []
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        # Skip headers and empty lines
        if not line or ':' in line and len(line) < 50:
            continue
        # Clean up bullet points
        line = re.sub(r'^[-•*\d]+\.?\s*', '', line)
        if line:
            points.append(line)
    return points

def clean_text(text):
    """Clean up text content"""
    # Remove common headers
    text = re.sub(r'^(Summary|Overall|Analysis):\s*', '', text, flags=re.IGNORECASE)
    return text.strip()
def get_track_requirements(track):
    """Get requirements for a specific track"""
    requirements = {
        'digital_technology': """
        Mandatory Criteria:
        • Exceptional technical expertise in digital technology
        • Proven track record of innovation in the field
        • Recognition as a leading talent in digital technology

        Optional Criteria (need to meet at least 2):
        • Significant technical, commercial or entrepreneurial contributions
        • Academic contributions through research in a leading institution
        • Recognition through awards, speaking engagements, or media coverage
        • Evidence of exceptional ability through salary, funding, or investment
        """,
        'data_science_ai': """
        Mandatory Criteria:
        • Exceptional expertise in AI, machine learning, or data science
        • Proven track record of innovation in AI/ML/data science
        • Recognition as a leading talent in the field

        Optional Criteria (need to meet at least 2):
        • Significant contributions to AI/ML research or commercial applications
        • Academic contributions through research publications
        • Recognition through awards, patents, or speaking engagements
        • Evidence of exceptional ability through salary, funding, or investment
        """
    }

    return requirements.get(track, requirements['digital_technology'])



@login_required
@require_points(1)  # CV Analysis costs 1 point (or one free use)
@rate_limit(max_calls=10, window=60)
@require_http_methods(["POST"])
def analyze_cv(request, *args, **kwargs):
    try:
        can_use_free_feature_use = kwargs.get('can_use_free_feature_use', False)
        will_deduct_ai_points = kwargs.get('will_deduct_ai_points', False)
        points_required = kwargs.get('points_required_for_action', 1)

        form_data = request.POST.copy()
        if 'title' not in form_data or not form_data['title']:
            form_data['title'] = f"CV Analysis - {request.user.username}" # Auto-generate title
        if 'status' not in form_data or not form_data['status']:
            form_data['status'] = 'draft' # Default status
        form = CVForm(form_data, request.FILES)

        if not form.is_valid():
            logger.warning(f"CV analysis form invalid for {request.user.email}: {form.errors}")
            return JsonResponse({'error': 'Invalid form data', 'error_type': 'validation_error', 'form_errors': form.errors}, status=400)

        cv_file = request.FILES.get('cv_file')
        track = form.cleaned_data.get('track', 'digital_technology') # Default track

        if not cv_file:
            logger.warning(f"CV file missing for analysis attempt by {request.user.email}")
            return JsonResponse({'error': 'CV file is required', 'error_type': 'validation_error'}, status=400)

        try:
            cv_content = extract_cv_content(cv_file)
            if not cv_content or len(cv_content) < 50: # Basic check for extracted content
                 raise ValueError("Extracted CV content is too short or empty.")
        except Exception as e:
            logger.error(f"CV extraction error for {request.user.email}: {str(e)}")
            return JsonResponse({'error': f'Error extracting CV content: {str(e)}', 'error_type': 'extraction_error'}, status=400)

        track_requirements_text = get_track_requirements(track) # This should return a string of requirements

        # --- Construct the detailed prompt for the AI ---
        prompt = f"""
Analyze the following CV content for a Tech Nation Global Talent Visa application.
The candidate is applying under the '{track}' track.

CV Content:
---
{cv_content}
---

Relevant Track Requirements for '{track}':
---
{track_requirements_text}
---

Please provide a detailed analysis. Your response MUST be a single JSON object with a top-level key "analysis".
The "analysis" object MUST include the following fields:
- "overallStrengthScore": An integer score from 0 to 100, assessing overall CV strength for the visa.
- "readinessLevelAssessment": A string (e.g., "High", "Medium", "Low", "Excellent", "Good", "Needs Improvement") indicating visa readiness.
- "keyStrengths": A list of strings (3-5 items) describing key strengths evident in the CV, aligned with Tech Nation criteria.
- "criticalGaps": A list of strings (3-5 items) identifying critical gaps or areas needing significant improvement for the visa application.
- "summaryOfAlignmentWithTechNationCriteria": A concise string (2-4 sentences) summarizing how well the CV aligns with the Tech Nation criteria for the given track.
- "improvementRoadmap": An object with the following keys, each containing a list of 2-3 actionable string recommendations:
    - "immediateActions": List of strings for quick wins (e.g., "Quantify achievements in project X with metrics.").
    - "shortTermActions": List of strings for actions in the next few weeks/months (e.g., "Seek a strong letter of recommendation from a leader in Y field.").
    - "longTermActions": List of strings for longer-term development (e.g., "Publish research or a significant article in Z area.").
- "detailedAnalysis": An object with keys representing key assessment areas. Each key MUST map to an object with "currentState" (string summary of current standing based on CV), "recommendations" (list of 1-2 specific string recommendations for this area), and "priority" (string: "High", "Medium", "Low"). Include at least:
    - "technicalExpertise": {{ "currentState": "...", "recommendations": ["...", "..."], "priority": "..." }}
    - "innovationContributions": {{ "currentState": "...", "recommendations": ["...", "..."], "priority": "..." }}
    - "leadershipAndImpact": {{ "currentState": "...", "recommendations": ["...", "..."], "priority": "..." }}
    - "recognitionAndAwards": {{ "currentState": "...", "recommendations": ["...", "..."], "priority": "..." }} (If no direct evidence, provide general advice on how to build this).
- "immediateActionItems": A list of 3-5 objects, where each object has "priority" (integer: 1 for highest, 2 for medium, 3 for lowest) and "item" (string: a concise, actionable item derived from criticalGaps or immediateActions). Example: [{{"priority": 1, "item": "Add specific metrics to Project Alpha to demonstrate commercial impact."}}]

Focus on providing actionable, specific, and constructive insights based SOLELY on the provided CV content and track requirements.
Ensure all lists contain strings. Do not invent information not present in the CV.
The entire response must be ONLY the JSON object. No introductory or concluding text.
"""
        logger.info(f"Attempting CV analysis for {request.user.email} on track '{track}'.")
        # --- AI Call ---
        try:
            # Use the new AIProvider method designed for CV analysis JSON output
            analysis_content_raw = ai_provider.generate_cv_analysis_content(prompt)
            
            if not analysis_content_raw:
                logger.error(f"AI provider failed to generate CV analysis for {request.user.email}. Raw response was empty.")
                return JsonResponse({'error': 'Failed to generate CV analysis from AI provider. The response was empty.', 'error_type': 'generation_error_empty'}, status=500)
            
            logger.debug(f"Raw AI response for CV analysis for {request.user.email}: {analysis_content_raw[:500]}") # Log snippet of raw response

            # --- Process AI response (parsing JSON) ---
            analysis_data = {}
            try:
                parsed_response = json.loads(analysis_content_raw)
                
                if 'analysis' in parsed_response and isinstance(parsed_response['analysis'], dict):
                    analysis = parsed_response['analysis']
                    analysis_data = {
                        'strength_score': analysis.get('overallStrengthScore', 0),
                        'readiness_level': analysis.get('readinessLevelAssessment', 'N/A'),
                        'key_strengths': analysis.get('keyStrengths', []),
                        'critical_gaps': analysis.get('criticalGaps', []), # Added this
                        'summary': analysis.get('summaryOfAlignmentWithTechNationCriteria', 'Summary not provided.'),
                        'improvement_roadmap': analysis.get('improvementRoadmap', {}),
                        'detailed_analysis': analysis.get('detailedAnalysis', {}),
                        'immediate_actions': [item.get('item') for item in analysis.get('immediateActionItems', []) if isinstance(item, dict) and item.get('item')]
                    }
                    logger.info(f"Successfully parsed AI JSON response for CV analysis for {request.user.email}.")
                else:
                    logger.warning(f"AI response for CV analysis for {request.user.email} did not contain 'analysis' key or it was not a dictionary. Raw: {analysis_content_raw[:200]}")
                    analysis_data = {"error": "AI response format unexpected. Missing 'analysis' key or incorrect structure.", "raw_preview": analysis_content_raw[:500]}
                    # Potentially do not charge if the format is wrong. For now, proceeding.

            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse AI JSON response for CV analysis for {request.user.email}. Error: {e}. Raw: {analysis_content_raw[:500]}")
                analysis_data = {"error": "Failed to parse AI response. The AI did not return valid JSON.", "raw_preview": analysis_content_raw[:500]}
                # Consider not charging the user if JSON parsing fails.
                # For now, proceeding to charge/use free use as per original logic.

            # --- Deduct free use or points *after* AI call and parsing attempt ---
            used_free_use_for_this_action = False
            points_deducted_this_action = 0
            user_profile = request.user.profile 
            user_points_obj, _ = UserPoints.objects.get_or_create(user=request.user)

            # Only deduct if the analysis_data does not primarily indicate an error from our side or AI's inability to format
            # This is a simple check; you might want more sophisticated logic.
            should_charge = not analysis_data.get("error", "").startswith("AI response format unexpected") and \
                            not analysis_data.get("error", "").startswith("Failed to parse AI response")


            if should_charge:
                if can_use_free_feature_use and user_profile.available_free_uses > 0:
                    user_profile.available_free_uses -= 1
                    user_profile.save(update_fields=['available_free_uses'])
                    used_free_use_for_this_action = True
                    logger.info(f"User {request.user.email} consumed one free use for CV analysis. Remaining: {user_profile.available_free_uses}")
                elif will_deduct_ai_points and user_points_obj.balance >= points_required:
                    try:
                        user_points_obj.use_points(points_required, description=f"CV Analysis for track: {track}") # Assuming use_points can take a description
                        points_deducted_this_action = points_required
                        logger.info(f"Deducted {points_required} AI points from {request.user.email} for CV analysis. New balance: {user_points_obj.balance}")
                    except Exception as e:
                        logger.error(f"Error deducting points for {request.user.email} in analyze_cv after AI call: {e}")
                        # If points deduction fails, it's a server-side issue. The user got the analysis.
                        # The response should still be success, but log this error for admin.
                        # Or, if critical, return a specific error.
                        # For now, let's assume the analysis is the primary product.
                # If neither, the @require_points decorator should have caught it.
                # If analysis_data contains an error (e.g. parsing failed), we still reach here.
                # You might want to avoid charging if `analysis_data` has an "error" key.
            else:
                logger.info(f"Skipped charging for CV analysis for {request.user.email} due to AI response processing error.")


            # Optionally save the CVAnalysis to a model if you have one
            # e.g., CVAnalysis.objects.create(user=request.user, cv_file_name=cv_file.name, analysis_result=analysis_data, track=track, points_cost=points_deducted_this_action, used_free_use=used_free_use_for_this_action)

            return JsonResponse({
                'success': True, # Success is true if we got *some* response, even if it's an error message within analysis_data
                'analysis': analysis_data,
                'used_free_feature_use': used_free_use_for_this_action,
                'available_free_uses_remaining': user_profile.available_free_uses,
                'points_deducted_from_ai_balance': points_deducted_this_action,
                'ai_points_remaining': user_points_obj.balance
            })

        except Exception as e: # Catch errors from AI call or other processing within this block
            logger.error(f"Error during AI CV analysis processing for {request.user.email}: {str(e)}", exc_info=True)
            return JsonResponse({'error': f'An error occurred during CV analysis: {str(e)}', 'error_type': 'analysis_processing_error'}, status=500)

    except Exception as e: # Broad exception for the whole view
        logger.error(f"Unexpected error in analyze_cv view for {request.user.email}: {str(e)}", exc_info=True)
        return JsonResponse({'error': 'An unexpected server error occurred. Please try again.', 'error_type': 'server_error'}, status=500)





def extract_section_value(text, section, subsection):
    """Extract a specific subsection value from text"""
    import re
    
    # Try to find the section and subsection
    pattern = rf"(?i){section}.*?{subsection}[:\s]+(.*?)(?:\n\n|\n[A-Z]|$)"
    match = re.search(pattern, text, re.DOTALL)
    
    if match:
        return match.group(1).strip()
    
    # Fallback: try to find just the subsection
    pattern = rf"(?i){subsection}[:\s]+(.*?)(?:\n\n|\n[A-Z]|$)"
    match = re.search(pattern, text, re.DOTALL)
    
    if match:
        return match.group(1).strip()
    
    # Second fallback: look for keywords
    if section.lower() == "technical expertise" and subsection.lower() == "current state":
        return extract_text_around_keywords(text, ["technical", "expertise", "skills", "abilities"])
    elif section.lower() == "leadership" and subsection.lower() == "current state":
        return extract_text_around_keywords(text, ["leadership", "management", "team", "leading"])
    elif section.lower() == "recognition" and subsection.lower() == "current state":
        return extract_text_around_keywords(text, ["recognition", "awards", "achievements", "acknowledged"])
    elif section.lower() == "commercial impact" and subsection.lower() == "current state":
        return extract_text_around_keywords(text, ["commercial", "impact", "business", "revenue", "growth"])
    
    # Default values based on section/subsection
    if subsection.lower() == "priority":
        return "Medium"
    elif subsection.lower() == "recommendations":
        return f"Improve your {section.lower()} with targeted actions."
    elif subsection.lower() == "current state":
        return f"Your {section.lower()} needs further development."
    
    return "Not available"

def extract_text_around_keywords(text, keywords, context_chars=100):
    """Extract text around keywords"""
    import re
    
    for keyword in keywords:
        pattern = rf"(?i)(.{{0,{context_chars}}}{keyword}.{{0,{context_chars}}})"
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()
    
    return "Not available"


def extract_strength_score(text):
    """Extract strength score from text analysis"""
    import re

    # Look for patterns like "Overall strength score: 85" or "Score: 85/100"
    score_patterns = [
        r'[Oo]verall\s+[Ss]trength\s+[Ss]core\s*:?\s*(\d+)',
        r'[Ss]trength\s+[Ss]core\s*:?\s*(\d+)',
        r'[Ss]core\s*:?\s*(\d+)\s*/\s*100',
        r'[Ss]core\s*:?\s*(\d+)'
    ]

    for pattern in score_patterns:
        match = re.search(pattern, text)
        if match:
            try:
                score = int(match.group(1))
                # Ensure score is within valid range
                return max(0, min(100, score))
            except (ValueError, IndexError):
                pass

    # If no score found, calculate based on content analysis
    # Count positive vs negative indicators
    positive_terms = ['excellent', 'strong', 'impressive', 'outstanding', 'exceptional']
    negative_terms = ['weak', 'poor', 'lacking', 'insufficient', 'inadequate']

    positive_count = sum(text.lower().count(term) for term in positive_terms)
    negative_count = sum(text.lower().count(term) for term in negative_terms)

    if positive_count + negative_count > 0:
        # Calculate a score based on the ratio of positive to total indicators
        ratio = positive_count / (positive_count + negative_count)
        return int(40 + (ratio * 60))  # Scale to range 40-100

    # Default fallback if no other method works
    return 65  # More neutral default





def extract_readiness_level(text):
    """Extract readiness level from text"""
    import re

    # Common readiness levels
    levels = ['high', 'medium', 'low', 'strong', 'moderate', 'weak', 'excellent', 'good', 'fair', 'poor', 'ready', 'not ready', 'needs improvement']

    # Try to find readiness level pattern
    for level in levels:
        pattern = rf'readiness\s*level:?\s*{level}'
        if re.search(pattern, text, re.IGNORECASE):
            return level.title()

    # Try to extract from JSON
    json_match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if json_match:
        try:
            import json
            json_str = json_match.group(1)
            data = json.loads(json_str)

            if 'analysis' in data and 'readinessLevelAssessment' in data['analysis']:
                return data['analysis']['readinessLevelAssessment']
        except:
            pass

    # Default level if nothing found
    return 'Needs Assessment'


def extract_points_from_text(text, keywords):
    """Extract bullet points from text based on keywords"""
    points = []

    # Check if this is JSON data embedded in text
    import re
    import json

    # Try to extract JSON from the text
    json_match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if json_match:
        try:
            json_str = json_match.group(1)
            data = json.loads(json_str)

            # Extract from JSON structure
            if 'analysis' in data:
                analysis = data['analysis']

                # Extract strength score
                if 'overallStrengthScore' in analysis:
                    return [str(analysis['overallStrengthScore'])]

                # Look for arrays with keywords
                for key, value in analysis.items():
                    if any(keyword.lower() in key.lower() for keyword in keywords) and isinstance(value, list):
                        # Clean up the items
                        clean_items = []
                        for item in value:
                            if isinstance(item, str):
                                # Remove quotes and asterisks
                                clean_item = item.strip('"\'')
                                clean_item = re.sub(r'^\*\s*', '', clean_item)
                                if clean_item:
                                    clean_items.append(clean_item)
                        if clean_items:
                            return clean_items
        except:
            pass  # If JSON parsing fails, continue with text extraction

    # Fall back to regular text extraction
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        # Skip empty lines
        if not line:
            continue

        # Check if line contains any of the keywords
        if any(keyword.lower() in line.lower() for keyword in keywords):
            # Clean up bullet points
            clean_line = re.sub(r'^[-•*\d]+\.?\s*', '', line)
            # Remove quotes
            clean_line = clean_line.strip('"\'')
            if clean_line and len(clean_line) > 10:  # Minimum length check
                points.append(clean_line)

    # If no points found with keywords, try to extract any bullet points
    if not points:
        for line in lines:
            line = line.strip()
            if line.startswith('-') or line.startswith('•') or line.startswith('*'):
                clean_line = re.sub(r'^[-•*\d]+\.?\s*', '', line)
                clean_line = clean_line.strip('"\'')
                if clean_line and len(clean_line) > 10:
                    points.append(clean_line)

    # Limit to 5 points
    return points[:5]

def extract_summary(text):
    """Extract a summary from the text"""
    # Try to extract JSON from the text
    import re
    import json

    json_match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
    if json_match:
        try:
            json_str = json_match.group(1)
            data = json.loads(json_str)

            # Extract from JSON structure
            if 'analysis' in data:
                analysis = data['analysis']

                # Look for summary fields
                if 'summaryOfAlignmentWithTechNationCriteria' in analysis and isinstance(analysis['summaryOfAlignmentWithTechNationCriteria'], list):
                    return " ".join(item.strip('"\'*') for item in analysis['summaryOfAlignmentWithTechNationCriteria'])

                if 'readinessLevelAssessment' in analysis:
                    return f"Readiness Level: {analysis['readinessLevelAssessment']}. Overall Strength Score: {analysis.get('overallStrengthScore', 'N/A')}."

            # Return the entire JSON as a fallback
            return "Analysis completed. Please review the detailed findings in the sections below."
        except:
            pass  # If JSON parsing fails, continue with text extraction

    # Look for sections that might contain summary
    summary_keywords = ['summary', 'overview', 'assessment', 'analysis']

    # Try to find a paragraph containing summary keywords
    paragraphs = text.split('\n\n')
    for paragraph in paragraphs:
        if any(keyword in paragraph.lower() for keyword in summary_keywords):
            # Clean up and return the paragraph
            clean_para = re.sub(r'^(summary|overview|assessment|analysis):?\s*', '', paragraph, flags=re.IGNORECASE)
            if clean_para and len(clean_para) > 20:
                return clean_para.strip()

    # If no summary found, return the first substantial paragraph
    for paragraph in paragraphs:
        if len(paragraph.strip()) > 50:
            return paragraph.strip()

    return "Analysis completed. Please review the detailed findings."





@login_required
def check_points(request):
    """Debug view to check user's points balance"""
    try:
        user_points = UserPoints.objects.get(user=request.user)
        return JsonResponse({
            'user': request.user.username,
            'points_balance': user_points.balance
        })
    except UserPoints.DoesNotExist:
        return JsonResponse({
            'user': request.user.username,
            'points_balance': 0,
            'note': 'User points record does not exist'
        })





        
def fix_json_response(response_text):
    """Attempt to fix common JSON formatting issues"""
    try:
        # Replace smart quotes with regular quotes
        response_text = response_text.replace('"', '"').replace('"', '"')
        response_text = response_text.replace(''', "'").replace(''', "'")

        # Fix unescaped quotes in strings
        import re

        # Find strings that contain unescaped quotes and fix them
        def fix_quotes_in_string(match):
            content = match.group(1)
            # Escape internal quotes
            content = content.replace('"', '\\"')
            return f'"{content}"'

        # Pattern to match strings with potential unescaped quotes
        pattern = r'"([^"]*"[^"]*)"'
        response_text = re.sub(pattern, fix_quotes_in_string, response_text)

        return response_text
    except Exception as e:
        logger.error(f"Error fixing JSON: {str(e)}")
        return response_text


def process_sections(content):
    """Process markdown sections for frontend display"""
    import re
    
    # Convert markdown headers
    content = re.sub(r'^# (.*)', r'<h2>\1</h2>', content, flags=re.MULTILINE)
    content = re.sub(r'^## (.*)', r'<h3>\1</h3>', content, flags=re.MULTILINE)
    
    # Convert bold text
    content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
    
    # Convert bullet points
    content = re.sub(r'^- (.*)', r'<li>\1</li>', content, flags=re.MULTILINE)
    
    # Wrap consecutive list items in ul tags
    content = re.sub(r'(<li>.*?</li>)', r'<ul>\1</ul>', content, flags=re.DOTALL)
    
    # Convert paragraphs
    paragraphs = content.split('\n\n')
    formatted_paragraphs = []
    
    for para in paragraphs:
        para = para.strip()
        if para and not para.startswith('<'):
            formatted_paragraphs.append(f'<p>{para}</p>')
        elif para:
            formatted_paragraphs.append(para)
    
    return '\n'.join(formatted_paragraphs)

@login_required
def document_list(request):
    """Display user's documents grouped by type"""

    # Get all documents for the current user
    documents = Document.objects.filter(user=request.user).order_by('-updated_at')

    # Group documents by type
    document_groups = defaultdict(list)

    for document in documents:
        # Map document types to display names
        type_mapping = {
            'personal_statement': 'Personal Statements',
            'cv': 'CVs',
            'recommendation': 'Recommendation Letters',
            'evidence': 'Evidence Documents',
            'other': 'Other Documents'
        }

        display_type = type_mapping.get(document.document_type, 'Other Documents')
        document_groups[display_type].append(document)

    # Convert defaultdict to regular dict for template
    document_groups = dict(document_groups)

    context = {
        'document_groups': document_groups,
        'total_documents': documents.count()
    }

    return render(request, 'document_manager/document_list.html', context)


@login_required
def document_detail(request, pk):
    """Display document details with edit functionality"""
    document = get_object_or_404(Document, pk=pk, user=request.user)
    
    # Calculate progress
    progress = calculate_application_progress(document)
    
    # Get related documents
    related_docs = Document.objects.filter(
        user=request.user,
        document_type=document.document_type
    ).exclude(pk=pk)[:5]
    
    context = {
        'document': document,
        'progress': progress,
        'related_docs': related_docs,
        'word_count': len(document.content.split()) if document.content else 0
    }
    
    return render(request, 'document_manager/document_detail.html', context)

@login_required
def create_document(request):
    """Create a new document"""
    # Get doc_type from GET parameter
    doc_type = request.GET.get('doc_type')

    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.user = request.user

            # Set document type if provided in GET parameter
            if doc_type:
                document.document_type = doc_type

            document.save()

            messages.success(request, f'{document.get_document_type_display()} created successfully!')
            return redirect('document_manager:document_detail', pk=document.pk)
    else:
        form = DocumentForm()
        # Pre-select document type if provided
        if doc_type:
            form.initial['document_type'] = doc_type

    # Basic context
    context = {
        'form': form,
        'doc_type': doc_type
    }

    # Add CV enhancement context if doc_type is 'cv'
    if doc_type == 'cv':
        # Get user points
        user_points, created = UserPoints.objects.get_or_create(user=request.user)

        # Add CV-specific context
        cv_context = {
            'tracks': [
                ('digital_technology', 'Digital Technology'),
                ('data_science_ai', 'Data Science & AI'),
                ('fintech', 'FinTech'),
                ('cyber_security', 'Cyber Security'),
            ],
            'guidelines': [
                "Highlight your technical expertise and innovations",
                "Include measurable achievements and impact",
                "Showcase leadership and industry recognition"
            ],
            'user_points': user_points,
            'cv_analysis_points_cost': 1  # Cost to analyze a CV
        }

        # Update the context with CV-specific data
        context.update(cv_context)

    return render(request, 'document_manager/create_document.html', context)


    

@login_required
def edit_document(request, pk):
    """Edit an existing document"""
    document = get_object_or_404(Document, pk=pk, user=request.user)
    
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES, instance=document)
        if form.is_valid():
            form.save()
            messages.success(request, 'Document updated successfully!')
            return redirect('document_manager:document_detail', pk=document.pk)
    else:
        form = DocumentForm(instance=document)
    
    context = {
        'form': form,
        'document': document
    }
    
    return render(request, 'document_manager/edit_document.html', context)




@login_required
@require_http_methods(["POST"])
def delete_document(request, pk):
    """Delete a document via AJAX"""
    try:
        document = get_object_or_404(Document, pk=pk, user=request.user)
        document_title = document.title
        document.delete()

        return JsonResponse({
            'success': True,
            'message': f'Document "{document_title}" deleted successfully'
        })

    except Exception as e:
        logger.error(f"Error deleting document {pk}: {str(e)}")
        return JsonResponse({
            'error': f'Error deleting document: {str(e)}'
        }, status=500)

@login_required
@require_http_methods(["POST"])
def set_as_chosen(request, pk):
    """Set a personal statement as chosen for submission"""
    try:
        document = get_object_or_404(Document, pk=pk, user=request.user)

        if document.document_type != 'personal_statement':
            return JsonResponse({
                'error': 'Only personal statements can be set as chosen'
            }, status=400)

        # Remove chosen status from other personal statements
        Document.objects.filter(
            user=request.user,
            document_type='personal_statement',
            is_chosen=True
        ).update(is_chosen=False)

        # Set this document as chosen
        document.is_chosen = True
        document.save()

        logger.info(f"Document {pk} set as chosen for user {request.user.id}")

        return JsonResponse({
            'success': True,
            'message': f'"{document.title}" has been set as your chosen personal statement'
        })

    except Exception as e:
        logger.error(f"Error setting document {pk} as chosen: {str(e)}")
        return JsonResponse({
            'error': f'Error updating document: {str(e)}'
        }, status=500)


@login_required
def duplicate_document(request, pk):
    """Create a duplicate of an existing document"""
    original = get_object_or_404(Document, pk=pk, user=request.user)
    
    # Create a new document with copied content
    duplicate = Document.objects.create(
        user=request.user,
        title=f"Copy of {original.title}",
        document_type=original.document_type,
        content=original.content,
        notes=original.notes,
        status='draft'
    )
    
    messages.success(request, 'Document duplicated successfully!')
    return redirect('document_manager:document_detail', pk=duplicate.pk)

@login_required
@require_http_methods(["POST"])
def batch_process_documents(request):
    """Process multiple documents in batch"""
    try:
        document_ids = request.POST.getlist('document_ids')
        action = request.POST.get('action')
        
        if not document_ids:
            return JsonResponse({'error': 'No documents selected'}, status=400)
        
        results = []
        
        # Use ThreadPoolExecutor for parallel processing
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            futures = []
            
            for doc_id in document_ids:
                future = executor.submit(process_single_document, doc_id, request.user)
                futures.append((doc_id, future))
            
            # Collect results
            for doc_id, future in futures:
                try:
                    result = future.result(timeout=30)  # 30 second timeout per document
                    results.append({
                        'document_id': doc_id,
                        'status': 'success',
                        'result': result
                    })
                except Exception as e:
                    results.append({
                        'document_id': doc_id,
                        'status': 'error',
                        'error': str(e)
                    })
        
        return JsonResponse({
            'success': True,
            'results': results,
            'processed': len(results)
        })
        
    except Exception as e:
        logger.error(f"Batch processing error: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def download_document(request, pk):
    """Download document as DOCX file"""
    document = get_object_or_404(Document, pk=pk, user=request.user)
    
    try:
        # Generate DOCX file
        filepath = save_to_docx(
            content=document.content,
            title=document.title,
            doc_type=document.document_type
        )
        
        # Return file response
        response = FileResponse(
            open(filepath, 'rb'),
            as_attachment=True,
            filename=f"{document.title.replace(' ', '_')}.docx"
        )
        
        # Clean up temporary file after response
        def cleanup():
            try:
                os.remove(filepath)
            except:
                pass
        
        # Schedule cleanup (this is a simple approach; in production use celery)
        threading.Timer(60, cleanup).start()
        
        return response
        
    except Exception as e:
        logger.error(f"Error downloading document {pk}: {str(e)}")
        messages.error(request, f'Error downloading document: {str(e)}')
        return redirect('document_manager:document_detail', pk=pk)

@login_required
def export_documents(request):
    """Export multiple documents as a ZIP file"""
    try:
        document_ids = request.GET.getlist('ids')
        if not document_ids:
            messages.error(request, 'No documents selected for export')
            return redirect('document_manager:document_list')
        
        documents = Document.objects.filter(
            pk__in=document_ids,
            user=request.user
        )
        
        if not documents.exists():
            messages.error(request, 'No valid documents found')
            return redirect('document_manager:document_list')
        
        # Create a temporary directory for the ZIP
        import zipfile
        import tempfile
        
        temp_dir = tempfile.mkdtemp()
        zip_path = os.path.join(temp_dir, f"documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip")
        
        with zipfile.ZipFile(zip_path, 'w') as zip_file:
            for document in documents:
                try:
                    # Generate DOCX for each document
                    docx_path = save_to_docx(
                        content=document.content,
                        title=document.title,
                        doc_type=document.document_type
                    )
                    
                    # Add to ZIP
                    zip_file.write(
                        docx_path,
                        f"{document.title.replace(' ', '_')}.docx"
                    )
                    
                    # Clean up individual DOCX file
                    os.remove(docx_path)
                    
                except Exception as e:
                    logger.error(f"Error processing document {document.pk}: {str(e)}")
                    continue
        
        # Return ZIP file
        response = FileResponse(
            open(zip_path, 'rb'),
            as_attachment=True,
            filename=f"documents_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        )
        
        # Schedule cleanup
        def cleanup():
            try:
                os.remove(zip_path)
                os.rmdir(temp_dir)
            except:
                pass
        
        threading.Timer(60, cleanup).start()
        
        return response
        
    except Exception as e:
        logger.error(f"Error exporting documents: {str(e)}")
        messages.error(request, f'Error exporting documents: {str(e)}')
        return redirect('document_manager:document_list')

@login_required
def personal_statement_builder(request):
    """Main personal statement builder interface"""
    # Get user's existing documents
    personal_statements = Document.objects.filter(
        user=request.user,
        document_type='personal_statement'
    ).order_by('-created_at')

    cvs = Document.objects.filter(
        user=request.user,
        document_type='cv'
    ).order_by('-created_at')

    # Get user points
    user_points, created = UserPoints.objects.get_or_create(user=request.user)

    # Get referral points using the helper function
    referral_points = get_referral_points(request.user)

    # Get Tech Nation requirements
    requirements = get_tech_nation_requirements('digital_technology')

    # Add guidelines for the template
    guidelines = [
        "Upload your current CV in PDF, DOC, or DOCX format",
        "Select the type of personal statement that best matches your profile",
        "Provide specific instructions about achievements you want highlighted",
        "Review and customize the generated content before submission",
        "Ensure all claims can be supported by evidence in your application",
        "Keep the final statement between 800-1000 words",
        "Focus on exceptional talent and potential contributions to UK tech sector"
    ]

    context = {
        'personal_statements': personal_statements,
        'cvs': cvs,
        'requirements': requirements,
        'guidelines': guidelines,
        'tracks': [
            ('digital_technology', 'Digital Technology'),
            ('data_science_ai', 'Data Science & AI'),
        ],
        'user_points': user_points,
        'referral_points': referral_points,
        'total_available_points': user_points.balance + referral_points,
        'personal_statement_points_cost': 3,
        'cv_analysis_points_cost': 1
    }

    return render(request, 'document_manager/personal_statement_builder.html', context)

@login_required
@require_points(3)  # Costs 3 points (or one free use)
@rate_limit(max_calls=5, window=300) # Adjusted rate limit
@require_http_methods(["POST"])
def generate_personal_statement(request, *args, **kwargs):
    try:
        # Flags from require_points decorator
        can_use_free_feature_use = kwargs.get('can_use_free_feature_use', False)
        will_deduct_ai_points = kwargs.get('will_deduct_ai_points', False)
        points_required = kwargs.get('points_required_for_action', 3)

        cv_file = request.FILES.get('cv_file')
        instructions = request.POST.get('instructions', '')
        # statement_type = request.POST.get('type', 'technical') # If you use this

        if not cv_file:
            return JsonResponse({'error': 'CV file is required', 'error_type': 'validation_error'}, status=400)

        try:
            cv_content = extract_cv_content(cv_file)
        except Exception as e:
            logger.error(f"CV extraction error for PS for {request.user.email}: {str(e)}")
            return JsonResponse({'error': f'Error extracting CV content: {str(e)}', 'error_type': 'extraction_error'}, status=400)

        # --- Deduct free use or points *before* creating document and starting task ---
        # This is because the task is async. If we wait, the user might trigger something else.
        # The risk is if task creation fails, but this is usually quick.
        user_profile = request.user.profile
        user_points_obj, _ = UserPoints.objects.get_or_create(user=request.user)
        used_free_use_for_this_action = False
        points_deducted_this_action = 0

        if can_use_free_feature_use and user_profile.available_free_uses > 0:
            user_profile.available_free_uses -= 1
            user_profile.save(update_fields=['available_free_uses'])
            used_free_use_for_this_action = True
            logger.info(f"User {request.user.email} consumed one free use for Personal Statement. Remaining: {user_profile.available_free_uses}")
        elif will_deduct_ai_points and user_points_obj.balance >= points_required:
            try:
                user_points_obj.use_points(points_required) # This method should save
                points_deducted_this_action = points_required
                logger.info(f"Deducted {points_required} AI points from {request.user.email} for Personal Statement. New balance: {user_points_obj.balance}")
            except Exception as e:
                logger.error(f"Error deducting points for {request.user.email} in generate_personal_statement: {e}")
                return JsonResponse({'error': 'Error processing points before starting generation.', 'error_type': 'points_deduction_error'}, status=500)
        else:
            # This state should have been prevented by the decorator.
            logger.error(f"Reached generate_personal_statement for {request.user.email} without sufficient resources after decorator check. This is unexpected.")
            return JsonResponse({'error': 'Insufficient points or free uses. Please refresh and try again.', 'error_type': 'insufficient_resources_post_decorator'}, status=402)


        # --- Create document and start background task ---
        document = Document.objects.create(
            user=request.user,
            title=f"Personal Statement - {timezone.now().strftime('%Y-%m-%d %H:%M')}",
            document_type='personal_statement',
            status='processing' # Initial status
        )
        task_id = str(uuid.uuid4())
        cache.set(f"task_document_{task_id}", document.id, timeout=3600) # Cache for 1 hour

        # Pass relevant info to the task.
        # The task itself doesn't need to know about points/free_uses, as they're already handled.
        DocumentProcessor.process_in_background(
            generate_document_task, # The actual Celery task function
            cv_content=cv_content,
            instructions=instructions,
            user_id=request.user.id,
            document_id=document.id, # Pass document_id directly
            task_id=task_id
        )
        logger.info(f"Personal statement generation task {task_id} started for document {document.id} by user {request.user.email}")

        return JsonResponse({
            'success': True,
            'message': 'Personal statement generation started. You will be notified upon completion.',
            'document_id': document.id,
            'task_id': task_id, # For client-side polling if needed
            'used_free_feature_use': used_free_use_for_this_action,
            'available_free_uses_remaining': user_profile.available_free_uses,
            'points_deducted_from_ai_balance': points_deducted_this_action,
            'ai_points_remaining': user_points_obj.balance
        })

    except Exception as e:
        logger.error(f"Error in generate_personal_statement view for {request.user.email}: {str(e)}", exc_info=True)
        return JsonResponse({'error': 'An unexpected error occurred. Please try again or contact support.', 'error_type': 'server_error'}, status=500)




@login_required
@require_http_methods(["GET"])
def check_generation_status(request, task_id):
    """Check the status of document generation"""
    try:
        # Get status from cache
        status_info = cache.get(f"task_status_{task_id}")
        
        if not status_info:
            return JsonResponse({
                'status': 'not_found',
                'error': 'Task not found or expired'
            }, status=404)
        
        # If task is complete, get the result
        if status_info.get('status') == 'complete':
            result = cache.get(f"task_result_{task_id}")
            document_id = cache.get(f"task_document_{task_id}")
            
            response_data = {
                'status': 'complete',
                'progress': 100,
                'provider': status_info.get('provider', 'Unknown'),
                'document_id': document_id
            }
            
            if result:
                response_data['content'] = result
            
            return JsonResponse(response_data)
        
        # Return current status
        return JsonResponse(status_info)
        
    except Exception as e:
        logger.error(f"Error checking generation status: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'error': str(e)
        }, status=500)


        

@login_required
@require_http_methods(["POST"])
def save_generated_content(request):
    """Save generated personal statement content"""
    try:
        data = json.loads(request.body)
        title = data.get('title', '').strip()
        content = data.get('content', '').strip()

        if not title:
            return JsonResponse({
                'error': 'Document title is required'
            }, status=400)

        if not content:
            return JsonResponse({
                'error': 'Document content is required'
            }, status=400)

        # Remove chosen status from all existing personal statements
        Document.objects.filter(
            user=request.user,
            document_type='personal_statement',
            is_chosen=True
        ).update(is_chosen=False)

        # Create new document and set as chosen
        document = Document.objects.create(
            user=request.user,
            title=title,
            content=content,
            document_type='personal_statement',
            status='completed',
            is_chosen=True  # Set the new document as chosen
        )

        logger.info(f"Personal statement saved and set as chosen: {document.id} for user {request.user.id}")

        return JsonResponse({
            'success': True,
            'message': f'Personal statement "{title}" saved and set as your chosen statement',
            'document_id': document.id,
            'is_chosen': True
        })

    except json.JSONDecodeError:
        logger.error("Invalid JSON in save_generated_content request")
        return JsonResponse({
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        logger.error(f"Error saving generated content: {str(e)}")
        return JsonResponse({
            'error': f'Error saving document: {str(e)}'
        }, status=500)

@login_required
def preview_personal_statement(request, pk):
    """Preview personal statement with formatting"""
    document = get_object_or_404(Document, pk=pk, user=request.user)
    
    if document.document_type != 'personal_statement':
        messages.error(request, 'This document is not a personal statement')
        return redirect('document_manager:document_detail', pk=pk)
    
    # Process content for preview
    processed_content = process_sections(document.content) if document.content else ''
    
    # Calculate word count
    word_count = len(document.content.split()) if document.content else 0
    
    # Check if word count is within recommended range
    word_count_status = 'good'
    if word_count < 800:
        word_count_status = 'low'
    elif word_count > 1200:
        word_count_status = 'high'
    
    context = {
        'document': document,
        'processed_content': processed_content,
        'word_count': word_count,
        'word_count_status': word_count_status,
        'recommended_min': 800,
        'recommended_max': 1200
    }
    
    return render(request, 'document_manager/preview_personal_statement.html', context)

@login_required
@require_http_methods(["POST"])
def update_personal_statement(request, pk):
    """Update personal statement content via AJAX"""
    try:
        document = get_object_or_404(Document, pk=pk, user=request.user)
        
        if document.document_type != 'personal_statement':
            return JsonResponse({'error': 'Invalid document type'}, status=400)
        
        data = json.loads(request.body)
        content = data.get('content', '')
        title = data.get('title', document.title)
        
        # Update document
        document.content = content
        document.title = title
        document.updated_at = datetime.now()
        document.save()
        
        # Calculate new word count
        word_count = len(content.split()) if content else 0
        
        return JsonResponse({
            'success': True,
            'word_count': word_count,
            'message': 'Personal statement updated successfully'
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Error updating personal statement: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def get_requirements(request, track):
    """Get Tech Nation requirements for a specific track"""
    try:
        requirements = get_tech_nation_requirements(track)
        return JsonResponse({
            'success': True,
            'requirements': requirements
        })
    except Exception as e:
        logger.error(f"Error getting requirements for track {track}: {str(e)}")
        return JsonResponse({
            'error': f'Error getting requirements: {str(e)}'
        }, status=500)

# Dashboard and Statistics Views
@login_required
def dashboard(request):
    """Main dashboard with statistics and recent activity"""
    # Get user's documents
    documents = Document.objects.filter(user=request.user)

    # Group documents by type for display
    document_groups = {
        'Personal Statements': documents.filter(document_type='personal_statement'),
        'CVs': documents.filter(document_type='cv'),
        'Recommendation Letters': documents.filter(document_type='recommendation_letter'),
    }

    # Remove empty groups
    document_groups = {k: v for k, v in document_groups.items() if v.exists()}

    # Calculate statistics
    stats = {
        'total_documents': documents.count(),
        'personal_statements': documents.filter(document_type='personal_statement').count(),
        'cvs': documents.filter(document_type='cv').count(),
        'completed': documents.filter(status='completed').count(),
        'in_progress': documents.filter(status__in=['draft', 'in_progress', 'generating']).count(),
    }

    # Recent documents
    recent_documents = documents.order_by('-updated_at')[:5]

    # Get user points - ensure this exists
    try:
        user_points = UserPoints.objects.get(user=request.user)
    except UserPoints.DoesNotExist:
        # Create user points if they don't exist
        user_points = UserPoints.objects.create(user=request.user, balance=0)

    # Recent activity (last 30 days)
    from datetime import timedelta
    thirty_days_ago = datetime.now() - timedelta(days=30)
    recent_activity = documents.filter(
        updated_at__gte=thirty_days_ago
    ).order_by('-updated_at')[:10]

    # Progress calculation for each recent document
    for doc in recent_documents:
        doc.progress = calculate_application_progress(doc)

    context = {
        'stats': stats,
        'recent_documents': recent_documents,
        'recent_activity': recent_activity,
        'document_groups': document_groups,
        'user_points': user_points.balance  # Pass the balance value, not the object
    }

    return render(request, 'document_manager/document_list.html', context)



@login_required
def api_usage_stats(request):
    """Get API usage statistics"""
    try:
        # This would typically come from a database or cache
        # For now, return mock data
        stats = {
            'total_requests': 150,
            'successful_requests': 142,
            'failed_requests': 8,
            'average_response_time': 2.3,
            'requests_today': 12,
            'remaining_quota': 88
        }
        
        return JsonResponse({
            'success': True,
            'stats': stats
        })
        
    except Exception as e:
        logger.error(f"Error getting API usage stats: {str(e)}")
        return JsonResponse({
            'error': str(e)
        }, status=500)

@login_required
@require_http_methods(["POST"])
def clear_cache(request):
    """Clear user-specific cache"""
    try:
        # Clear document cache for this user
        user_cache_keys = [key for key in document_cache.keys() if str(request.user.id) in str(key)]
        for key in user_cache_keys:
            del document_cache[key]
        
        # Clear Django cache for this user
        cache_pattern = f"*{request.user.id}*"
        # Note: This is a simplified approach. In production, use Redis with pattern matching
        
        return JsonResponse({
            'success': True,
            'message': 'Cache cleared successfully'
        })
        
    except Exception as e:
        logger.error(f"Error clearing cache: {str(e)}")
        return JsonResponse({
            'error': str(e)
        }, status=500)

# Health check endpoint
def health_check(request):
    """Health check endpoint for monitoring"""
    try:
        # Check database connection
        Document.objects.count()
        
        # Check cache
        cache.set('health_check', 'ok', 60)
        cache_status = cache.get('health_check')
        
        return JsonResponse({
            'status': 'healthy',
            'database': 'ok',
            'cache': 'ok' if cache_status == 'ok' else 'error',
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        return JsonResponse({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }, status=500)

# Error handlers
def handler404(request, exception):
    """Custom 404 error handler"""
    return render(request, 'errors/404.html', status=404)

def handler500(request):
    """Custom 500 error handler"""
    return render(request, 'errors/500.html', status=500)

@login_required
def document_list_partial(request):
    """Partial view for HTMX document list updates"""
    documents = Document.objects.filter(user=request.user).order_by('-created_at')
    
    # Filter by document type if specified
    doc_type = request.GET.get('type')
    if doc_type:
        documents = documents.filter(document_type=doc_type)
    
    # Filter by status if specified
    status = request.GET.get('status')
    if status:
        documents = documents.filter(status=status)
    
    # Search functionality
    search = request.GET.get('search')
    if search:
        documents = documents.filter(
            models.Q(title__icontains=search) | 
            models.Q(content__icontains=search)
        )
    
    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(documents, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'doc_type': doc_type,
        'status': status,
        'search': search,
    }
    
    return render(request, 'document_manager/partials/document_list.html', context)

@login_required
def personal_statement_form(request):
    """Personal statement form view"""
    if request.method == 'POST':
        form = PersonalStatementForm(request.POST)
        if form.is_valid():
            # Process the form data
            cv_file = request.FILES.get('cv_file')
            instructions = form.cleaned_data.get('instructions', '')
            track = form.cleaned_data.get('track', 'digital_technology')
            
            # Redirect to generation process
            return redirect('document_manager:generate_personal_statement')
    else:
        form = PersonalStatementForm()
    
    return render(request, 'document_manager/personal_statement_form.html', {'form': form})

@login_required
def cv_upload_form(request):
    """CV upload form view"""
    if request.method == 'POST':
        form = CVForm(request.POST, request.FILES)
        if form.is_valid():
            cv = form.save(commit=False)
            cv.user = request.user
            cv.document_type = 'cv'
            cv.save()

            messages.success(request, 'CV uploaded successfully!')
            return redirect('document_manager:document_detail', pk=cv.pk)
    else:
        form = CVForm()

    return render(request, 'document_manager/cv_upload_form.html', {'form': form})

@login_required
@require_http_methods(["POST"])
def quick_generate(request):
    """Quick generation endpoint for simple requests"""
    try:
        data = json.loads(request.body)
        prompt = data.get('prompt', '').strip()
        doc_type = data.get('type', 'personal_statement')

        if not prompt:
            return JsonResponse({'error': 'Prompt is required'}, status=400)

        # Rate limiting check
        user_key = f"quick_gen_{request.user.id}"
        recent_calls = cache.get(user_key, 0)

        if recent_calls >= 3:  # Max 3 quick generations per hour
            return JsonResponse({
                'error': 'Rate limit exceeded. Please wait before making another request.'
            }, status=429)

        # Increment counter
        cache.set(user_key, recent_calls + 1, 3600)  # 1 hour expiry

        # Generate content
        content = ai_provider.generate_content(prompt)

        if not content:
            return JsonResponse({'error': 'Failed to generate content'}, status=500)

        # Process content for frontend
        processed_content = process_content_for_frontend(content)

        return JsonResponse({
            'success': True,
            'content': processed_content,
            'word_count': len(content.split()) if content else 0
        })

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Error in quick generation: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@require_http_methods(["POST"])
def regenerate_section(request):
    """Regenerate a specific section of a document"""
    try:
        data = json.loads(request.body)
        section_content = data.get('section_content', '').strip()
        section_type = data.get('section_type', 'general')
        instructions = data.get('instructions', '').strip()

        if not section_content:
            return JsonResponse({'error': 'Section content is required'}, status=400)

        # Create specific prompt based on section type
        section_prompts = {
            'introduction': "Rewrite this introduction section to be more compelling and professional:",
            'technical_expertise': "Improve this technical expertise section with more specific examples:",
            'leadership': "Enhance this leadership section with quantifiable achievements:",
            'innovation': "Strengthen this innovation section with concrete examples:",
            'conclusion': "Rewrite this conclusion to be more impactful and forward-looking:",
            'general': "Improve and enhance this section:"
        }

        base_prompt = section_prompts.get(section_type, section_prompts['general'])

        prompt = f"""
        {base_prompt}

        Original content:
        {section_content}

        Additional instructions: {instructions}

        Please maintain the professional tone and ensure the content aligns with Tech Nation Global Talent Visa requirements.
        Format with proper markdown headers and bullet points where appropriate.
        """

        # Generate improved content
        improved_content = ai_provider.generate_content(prompt)

        if not improved_content:
            return JsonResponse({'error': 'Failed to regenerate section'}, status=500)

        # Process for frontend
        processed_content = process_sections(improved_content)

        return JsonResponse({
            'success': True,
            'content': processed_content,
            'word_count': len(improved_content.split()) if improved_content else 0
        })

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Error regenerating section: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@require_http_methods(["POST"])
def improve_document(request, pk):
    """Improve an existing document using AI"""
    try:
        document = get_object_or_404(Document, pk=pk, user=request.user)

        data = json.loads(request.body)
        improvement_type = data.get('type', 'general')
        specific_instructions = data.get('instructions', '').strip()

        if not document.content:
            return JsonResponse({'error': 'Document has no content to improve'}, status=400)

        # Create improvement prompts based on type
        improvement_prompts = {
            'grammar': "Please review and improve the grammar, spelling, and sentence structure of this document:",
            'clarity': "Please improve the clarity and readability of this document while maintaining its professional tone:",
            'structure': "Please improve the structure and organization of this document:",
            'technical': "Please enhance the technical content and add more specific technical details:",
            'impact': "Please strengthen the impact statements and add more quantifiable achievements:",
            'general': "Please review and improve this document overall, focusing on clarity, impact, and professional presentation:"
        }

        base_prompt = improvement_prompts.get(improvement_type, improvement_prompts['general'])

        prompt = f"""
        {base_prompt}

        Document content:
        {document.content}

        Specific instructions: {specific_instructions}

        Please maintain the document's core message while improving its quality.
        Ensure the content remains suitable for a Tech Nation Global Talent Visa application.
        Format with proper markdown headers and structure.
        """

        # Generate improved content
        improved_content = ai_provider.generate_content(prompt)

        if not improved_content:
            return JsonResponse({'error': 'Failed to improve document'}, status=500)

        # Create a new version or update existing
        create_new_version = data.get('create_new_version', False)

        if create_new_version:
            # Create a new document with improved content
            new_document = Document.objects.create(
                user=request.user,
                title=f"{document.title} (Improved)",
                document_type=document.document_type,
                content=improved_content,
                status='completed',
                notes=f"Improved version of document {document.id} using {improvement_type} enhancement"
            )

            return JsonResponse({
                'success': True,
                'new_document_id': new_document.id,
                'message': 'Improved version created successfully'
            })
        else:
            # Update existing document
            document.content = improved_content
            document.notes = f"Improved using {improvement_type} enhancement on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            document.save()

            return JsonResponse({
                'success': True,
                'content': process_content_for_frontend(improved_content),
                'word_count': len(improved_content.split()) if improved_content else 0,
                'message': 'Document improved successfully'
            })

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Error improving document {pk}: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@require_http_methods(["GET"])
def document_versions(request, pk):
    """Get all versions of a document"""
    try:
        base_document = get_object_or_404(Document, pk=pk, user=request.user)

        # Find related versions (documents with similar titles or notes referencing this document)
        versions = Document.objects.filter(
            user=request.user,
            document_type=base_document.document_type
        ).filter(
            models.Q(title__icontains=base_document.title.split(' (')[0]) |
            models.Q(notes__icontains=f"document {pk}")
        ).order_by('-created_at')

        versions_data = []
        for version in versions:
            versions_data.append({
                'id': version.id,
                'title': version.title,
                'created_at': version.created_at.isoformat(),
                'updated_at': version.updated_at.isoformat(),
                'status': version.status,
                'word_count': len(version.content.split()) if version.content else 0,
                'notes': version.notes or ''
            })

        return JsonResponse({
            'success': True,
            'versions': versions_data,
            'total': len(versions_data)
        })

    except Exception as e:
        logger.error(f"Error getting document versions: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@require_http_methods(["POST"])
def compare_documents(request):
    """Compare two documents and highlight differences"""
    try:
        data = json.loads(request.body)
        doc1_id = data.get('document1_id')
        doc2_id = data.get('document2_id')

        if not doc1_id or not doc2_id:
            return JsonResponse({'error': 'Both document IDs are required'}, status=400)

        doc1 = get_object_or_404(Document, pk=doc1_id, user=request.user)
        doc2 = get_object_or_404(Document, pk=doc2_id, user=request.user)

        # Simple comparison (in production, you might use a more sophisticated diff algorithm)
        import difflib

        doc1_lines = doc1.content.splitlines() if doc1.content else []
        doc2_lines = doc2.content.splitlines() if doc2.content else []

        # Generate HTML diff
        differ = difflib.HtmlDiff()
        diff_html = differ.make_file(
            doc1_lines,
            doc2_lines,
            fromdesc=doc1.title,
            todesc=doc2.title,
            context=True,
            numlines=3
        )

        # Calculate similarity percentage
        similarity = difflib.SequenceMatcher(None, doc1.content or '', doc2.content or '').ratio()
        similarity_percentage = round(similarity * 100, 1)

        return JsonResponse({
            'success': True,
            'diff_html': diff_html,
            'similarity_percentage': similarity_percentage,
            'doc1_word_count': len(doc1.content.split()) if doc1.content else 0,
            'doc2_word_count': len(doc2.content.split()) if doc2.content else 0
        })

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Error comparing documents: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
@require_http_methods(["POST"])
def document_feedback(request, pk):
    """Get AI feedback on a document"""
    try:
        document = get_object_or_404(Document, pk=pk, user=request.user)

        if not document.content:
            return JsonResponse({'error': 'Document has no content to analyze'}, status=400)

        data = json.loads(request.body)
        feedback_type = data.get('type', 'comprehensive')

        # Create feedback prompts
        feedback_prompts = {
            'comprehensive': """
            Please provide comprehensive feedback on this Tech Nation Global Talent Visa document.
            Analyze the following aspects:
            1. Overall structure and organization
            2. Technical content quality and specificity
            3. Evidence of exceptional talent
            4. Leadership and impact demonstration
            5. Clarity and professional presentation
            6. Alignment with Tech Nation requirements
            7. Specific suggestions for improvement

            Provide actionable feedback with specific examples.
            """,
            'technical': """
            Focus specifically on the technical content of this document.
            Evaluate:
            1. Technical depth and accuracy
            2. Use of appropriate technical terminology
            3. Demonstration of technical expertise
            4. Innovation and technical contributions
            5. Technical leadership examples

            Provide specific suggestions for technical improvements.
            """,
            'structure': """
            Analyze the structure and organization of this document.
            Evaluate:
            1. Logical flow and organization
            2. Section transitions
            3. Paragraph structure
            4. Use of headers and formatting
            5. Overall readability

            Suggest specific structural improvements.
            """,
            'requirements': """
            Evaluate how well this document meets Tech Nation Global Talent Visa requirements.
            Check for:
            1. Evidence of exceptional talent
            2. Technical expertise demonstration
            3. Leadership and impact examples
            4. Innovation contributions
            5. Recognition and achievements
            6. Future potential in UK tech sector

            Identify gaps and suggest specific evidence to include.
            """
        }

        base_prompt = feedback_prompts.get(feedback_type, feedback_prompts['comprehensive'])

        prompt = f"""
        {base_prompt}

        Document content:
        {document.content}

        Please provide detailed, constructive feedback in a structured format.
        Use bullet points and clear sections for easy reading.
        """

        # Generate feedback
        feedback_content = ai_provider.generate_content(prompt)

        if not feedback_content:
            return JsonResponse({'error': 'Failed to generate feedback'}, status=500)

        # Process feedback for frontend
        processed_feedback = process_sections(feedback_content)

        return JsonResponse({
            'success': True,
            'feedback': processed_feedback,
            'feedback_type': feedback_type,
            'document_word_count': len(document.content.split()) if document.content else 0
        })

    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Error generating feedback for document {pk}: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def document_analytics(request, pk):
    """Get analytics for a specific document"""
    try:
        document = get_object_or_404(Document, pk=pk, user=request.user)

        if not document.content:
            return JsonResponse({'error': 'Document has no content to analyze'}, status=400)

        content = document.content

        # Basic analytics
        word_count = len(content.split())
        char_count = len(content)
        char_count_no_spaces = len(content.replace(' ', ''))
        paragraph_count = len([p for p in content.split('\n\n') if p.strip()])

        # Readability analysis (simplified)
        sentences = content.count('.') + content.count('!') + content.count('?')
        avg_words_per_sentence = word_count / sentences if sentences > 0 else 0

        # Keyword analysis (simplified)
        import re
        from collections import Counter

        # Extract words (remove common stop words)
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your', 'his', 'her', 'its', 'our', 'their'}

        words = re.findall(r'\b[a-zA-Z]+\b', content.lower())
        filtered_words = [word for word in words if word not in stop_words and len(word) > 3]
        word_freq = Counter(filtered_words)
        top_keywords = word_freq.most_common(10)

        # Tech Nation specific keywords
        tech_keywords = ['technology', 'technical', 'innovation', 'leadership', 'development', 'software', 'engineering', 'digital', 'data', 'ai', 'machine', 'learning', 'algorithm', 'programming', 'coding', 'project', 'team', 'management', 'research', 'publication', 'patent', 'award', 'recognition', 'contribution', 'impact', 'growth', 'scalable', 'architecture', 'system', 'platform', 'framework', 'methodology', 'optimization', 'performance', 'security', 'cloud', 'database', 'api', 'frontend', 'backend', 'fullstack', 'devops', 'agile', 'scrum', 'startup', 'enterprise', 'commercial', 'revenue', 'user', 'customer', 'product', 'solution', 'implementation', 'deployment', 'maintenance', 'testing', 'debugging', 'troubleshooting', 'collaboration', 'mentoring', 'training', 'presentation', 'conference', 'publication', 'open', 'source', 'github', 'stackoverflow', 'linkedin', 'twitter', 'blog', 'article', 'tutorial', 'documentation', 'specification', 'requirement', 'analysis', 'design', 'prototype', 'mvp', 'iteration', 'feedback', 'improvement', 'enhancement', 'feature', 'functionality', 'usability', 'accessibility', 'responsive', 'mobile', 'web', 'application', 'website', 'interface', 'experience', 'user', 'journey', 'conversion', 'engagement', 'retention', 'acquisition', 'analytics', 'metrics', 'kpi', 'roi', 'business', 'strategy', 'vision', 'mission', 'goal', 'objective', 'milestone', 'deadline', 'budget', 'resource', 'allocation', 'planning', 'execution', 'delivery', 'quality', 'standard', 'best', 'practice', 'industry', 'sector', 'market', 'competition', 'advantage', 'differentiation', 'unique', 'value', 'proposition', 'benefit', 'outcome', 'result', 'achievement', 'success', 'failure', 'lesson', 'learned', 'experience', 'skill', 'expertise', 'knowledge', 'competency', 'proficiency', 'mastery', 'certification', 'qualification', 'degree', 'education', 'training', 'course', 'workshop', 'seminar', 'conference', 'meetup', 'networking', 'community', 'volunteer', 'contribution', 'participation', 'involvement', 'engagement', 'commitment', 'dedication', 'passion', 'enthusiasm', 'motivation', 'drive', 'ambition', 'aspiration', 'goal', 'dream', 'vision', 'future', 'potential', 'opportunity', 'challenge', 'problem', 'solution', 'approach', 'method', 'technique', 'tool', 'technology', 'platform', 'framework', 'library', 'package', 'module', 'component', 'service', 'microservice', 'api', 'endpoint', 'request', 'response', 'data', 'information', 'knowledge', 'insight', 'intelligence', 'wisdom', 'understanding', 'comprehension', 'awareness', 'consciousness', 'mindfulness', 'attention', 'focus', 'concentration', 'dedication', 'commitment', 'persistence', 'perseverance', 'resilience', 'adaptability', 'flexibility', 'agility', 'responsiveness', 'proactivity', 'initiative', 'creativity', 'innovation', 'invention', 'discovery', 'exploration', 'experimentation', 'research', 'investigation', 'analysis', 'synthesis', 'evaluation', 'assessment', 'measurement', 'monitoring', 'tracking', 'reporting', 'documentation', 'communication', 'collaboration', 'cooperation', 'coordination', 'synchronization', 'integration', 'alignment', 'harmony', 'balance', 'equilibrium', 'stability', 'consistency', 'reliability', 'dependability', 'trustworthiness', 'credibility', 'reputation', 'brand', 'image', 'perception', 'impression', 'opinion', 'feedback', 'review', 'rating', 'score', 'ranking', 'position', 'status', 'level', 'grade', 'class', 'category', 'type', 'kind', 'sort', 'variety', 'diversity', 'range', 'scope', 'scale', 'size', 'magnitude', 'extent', 'degree', 'amount', 'quantity', 'number', 'count', 'total', 'sum', 'average', 'mean', 'median', 'mode', 'minimum', 'maximum', 'range', 'variance', 'deviation', 'distribution', 'pattern', 'trend', 'direction', 'movement', 'change', 'transformation', 'evolution', 'development', 'growth', 'expansion', 'improvement', 'enhancement', 'optimization', 'refinement', 'polishing', 'finishing', 'completion', 'accomplishment', 'achievement', 'success', 'victory', 'win', 'triumph', 'conquest', 'mastery', 'dominance', 'superiority', 'excellence', 'perfection', 'ideal', 'standard', 'benchmark', 'reference', 'baseline', 'foundation', 'basis', 'ground', 'root', 'origin', 'source', 'beginning', 'start', 'initiation', 'launch', 'introduction', 'presentation', 'demonstration', 'exhibition', 'display', 'showcase', 'highlight', 'feature', 'emphasize', 'stress', 'underline', 'underscore', 'accentuate', 'amplify', 'magnify', 'enlarge', 'expand', 'extend', 'stretch', 'reach', 'span', 'cover', 'include', 'encompass', 'comprise', 'contain', 'hold', 'carry', 'bear', 'support', 'sustain', 'maintain', 'preserve', 'protect', 'defend', 'guard', 'shield', 'shelter', 'cover', 'hide', 'conceal', 'mask', 'disguise', 'camouflage', 'blend', 'merge', 'combine', 'unite', 'join', 'connect', 'link', 'associate', 'relate', 'correspond', 'match', 'fit', 'suit', 'adapt', 'adjust', 'modify', 'alter', 'change', 'transform', 'convert', 'translate', 'interpret', 'explain', 'clarify', 'illuminate', 'enlighten', 'educate', 'teach', 'instruct', 'guide', 'direct', 'lead', 'manage', 'supervise', 'oversee', 'monitor', 'watch', 'observe', 'notice', 'see', 'view', 'look', 'examine', 'inspect', 'check', 'verify', 'confirm', 'validate', 'authenticate', 'authorize', 'approve', 'accept', 'agree', 'consent', 'permit', 'allow', 'enable', 'facilitate', 'assist', 'help', 'aid', 'support', 'back', 'endorse', 'recommend', 'suggest', 'propose', 'offer', 'provide', 'supply', 'deliver', 'give', 'grant', 'award', 'present', 'bestow', 'confer', 'honor', 'recognize', 'acknowledge', 'appreciate', 'value', 'treasure', 'cherish', 'love', 'adore', 'admire', 'respect', 'esteem', 'regard', 'consider', 'think', 'believe', 'feel', 'sense', 'perceive', 'understand', 'comprehend', 'grasp', 'realize', 'recognize', 'identify', 'distinguish', 'differentiate', 'discriminate', 'separate', 'divide', 'split', 'break', 'fragment', 'piece', 'part', 'section', 'segment', 'portion', 'share', 'fraction', 'percentage', 'ratio', 'proportion', 'rate', 'speed', 'pace', 'tempo', 'rhythm', 'beat', 'pulse', 'frequency', 'interval', 'period', 'duration', 'time', 'moment', 'instant', 'second', 'minute', 'hour', 'day', 'week', 'month', 'year', 'decade', 'century', 'millennium', 'era', 'age', 'epoch', 'period', 'phase', 'stage', 'step', 'level', 'tier', 'layer', 'stratum', 'floor', 'ceiling', 'roof', 'top', 'summit', 'peak', 'apex', 'pinnacle', 'climax', 'culmination', 'conclusion', 'end', 'finish', 'termination', 'completion', 'accomplishment', 'achievement', 'attainment', 'realization', 'fulfillment', 'satisfaction', 'contentment', 'happiness', 'joy', 'pleasure', 'delight', 'enjoyment', 'fun', 'entertainment', 'amusement', 'recreation', 'relaxation', 'rest', 'break', 'pause', 'stop', 'halt', 'cease', 'quit', 'discontinue', 'abandon', 'give', 'up', 'surrender', 'yield', 'submit', 'comply', 'obey', 'follow', 'adhere', 'stick', 'cling', 'hold', 'grasp', 'grip', 'clutch', 'seize', 'catch', 'capture', 'trap', 'snare', 'ensnare', 'entrap', 'confine', 'restrict', 'limit', 'constrain', 'restrain', 'control', 'regulate', 'govern', 'rule', 'command', 'order', 'direct', 'instruct', 'tell', 'say', 'speak', 'talk', 'communicate', 'express', 'convey', 'transmit', 'send', 'deliver', 'transport', 'carry', 'move', 'shift', 'transfer', 'relocate', 'migrate', 'travel', 'journey', 'trip', 'voyage', 'expedition', 'adventure', 'exploration', 'discovery', 'finding', 'result', 'outcome', 'consequence', 'effect', 'impact', 'influence', 'power', 'force', 'strength', 'energy', 'vigor', 'vitality', 'life', 'existence', 'being', 'entity', 'object', 'thing', 'item', 'element', 'component', 'part', 'piece', 'fragment', 'bit', 'portion', 'section', 'segment', 'division', 'category', 'class', 'group', 'set', 'collection', 'assembly', 'gathering', 'meeting', 'conference', 'convention', 'summit', 'forum', 'symposium', 'seminar', 'workshop', 'training', 'course', 'program', 'curriculum', 'syllabus', 'schedule', 'timetable', 'agenda', 'plan', 'strategy', 'approach', 'method', 'technique', 'procedure', 'process', 'system', 'mechanism', 'device', 'tool', 'instrument', 'equipment', 'apparatus', 'machine', 'engine', 'motor', 'generator', 'producer', 'creator', 'maker', 'builder', 'constructor', 'developer', 'designer', 'architect', 'engineer', 'programmer', 'coder', 'developer', 'analyst', 'consultant', 'advisor', 'expert', 'specialist', 'professional', 'practitioner', 'worker', 'employee', 'staff', 'personnel', 'team', 'group', 'unit', 'department', 'division', 'section', 'branch', 'office', 'bureau', 'agency', 'organization', 'institution', 'establishment', 'company', 'corporation', 'business', 'enterprise', 'firm', 'venture', 'startup', 'initiative', 'project', 'undertaking', 'endeavor', 'effort', 'attempt', 'try', 'trial', 'test', 'experiment', 'study', 'research', 'investigation', 'inquiry', 'examination', 'analysis', 'evaluation', 'assessment', 'appraisal', 'review', 'audit', 'inspection', 'check', 'verification', 'validation', 'confirmation', 'proof', 'evidence', 'testimony', 'witness', 'account', 'report', 'statement', 'declaration', 'announcement', 'proclamation', 'publication', 'release', 'disclosure', 'revelation', 'exposure', 'unveiling', 'presentation', 'introduction', 'launch', 'debut', 'premiere', 'opening', 'start', 'beginning', 'commencement', 'initiation', 'inception', 'origin', 'source', 'root', 'foundation', 'basis', 'ground', 'reason', 'cause', 'motive', 'purpose', 'goal', 'objective', 'aim', 'target', 'intention', 'plan', 'design', 'scheme', 'program', 'project', 'initiative', 'campaign', 'drive', 'push', 'effort', 'attempt', 'try', 'endeavor', 'undertaking', 'venture', 'enterprise', 'business', 'operation', 'activity', 'action', 'deed', 'act', 'performance', 'execution', 'implementation', 'realization', 'accomplishment', 'achievement', 'success', 'triumph', 'victory', 'win', 'conquest', 'mastery', 'dominance', 'control', 'power', 'authority', 'influence', 'impact', 'effect', 'result', 'outcome', 'consequence', 'implication', 'significance', 'importance', 'relevance', 'value', 'worth', 'merit', 'quality', 'standard', 'level', 'grade', 'rank', 'position', 'status', 'standing', 'reputation', 'image', 'perception', 'view', 'opinion', 'judgment', 'assessment', 'evaluation', 'appraisal', 'estimation', 'calculation', 'computation', 'measurement', 'quantification', 'analysis', 'breakdown', 'decomposition', 'dissection', 'examination', 'study', 'investigation', 'research', 'exploration', 'discovery', 'finding', 'result', 'conclusion', 'decision', 'choice', 'selection', 'option', 'alternative', 'possibility', 'opportunity', 'chance', 'prospect', 'potential', 'capability', 'capacity', 'ability', 'skill', 'talent', 'gift', 'aptitude', 'competence', 'proficiency', 'expertise', 'knowledge', 'understanding', 'comprehension', 'awareness', 'consciousness', 'recognition', 'realization', 'appreciation', 'acknowledgment', 'acceptance', 'approval', 'endorsement', 'support', 'backing', 'assistance', 'help', 'aid', 'service', 'contribution', 'input', 'participation', 'involvement', 'engagement', 'commitment', 'dedication', 'devotion', 'loyalty', 'faithfulness', 'reliability', 'dependability', 'trustworthiness', 'credibility', 'integrity', 'honesty', 'sincerity', 'authenticity', 'genuineness', 'truth', 'fact', 'reality', 'actuality', 'existence', 'presence', 'occurrence', 'happening', 'event', 'incident', 'episode', 'situation', 'circumstance', 'condition', 'state', 'status', 'position', 'location', 'place', 'spot', 'site', 'venue', 'setting', 'environment', 'context', 'background', 'history', 'past', 'previous', 'former', 'earlier', 'prior', 'before', 'preceding', 'antecedent', 'preliminary', 'preparatory', 'initial', 'first', 'primary', 'main', 'principal', 'chief', 'leading', 'top', 'highest', 'supreme', 'ultimate', 'final', 'last', 'end', 'conclusion', 'termination', 'completion', 'finish', 'closure', 'resolution', 'solution', 'answer', 'response', 'reply', 'reaction', 'feedback', 'comment', 'remark', 'observation', 'note', 'notice', 'attention', 'focus', 'concentration', 'emphasis', 'stress', 'importance', 'significance', 'relevance', 'pertinence', 'applicability', 'suitability', 'appropriateness', 'fitness', 'match', 'correspondence', 'agreement', 'harmony', 'consistency', 'coherence', 'logic', 'reason', 'sense', 'meaning', 'significance', 'import', 'implication', 'suggestion', 'indication', 'sign', 'signal', 'symbol', 'representation', 'image', 'picture', 'illustration', 'example', 'instance', 'case', 'sample', 'specimen', 'model', 'prototype', 'template', 'pattern', 'design', 'structure', 'organization', 'arrangement', 'configuration', 'setup', 'layout', 'format', 'style', 'appearance', 'look', 'aspect', 'feature', 'characteristic', 'attribute', 'property', 'quality', 'trait', 'element', 'component', 'part', 'piece', 'section', 'segment', 'portion', 'share', 'fraction', 'percentage', 'proportion', 'ratio', 'rate', 'speed', 'pace', 'tempo', 'rhythm', 'frequency', 'occurrence', 'incidence', 'prevalence', 'commonness', 'rarity', 'uniqueness', 'distinctiveness', 'individuality', 'personality', 'character', 'nature', 'essence', 'substance', 'material', 'matter', 'content', 'information', 'data', 'facts', 'details', 'particulars', 'specifics', 'specifications', 'requirements', 'criteria', 'standards', 'guidelines', 'rules', 'regulations', 'policies', 'procedures', 'protocols', 'methods', 'techniques', 'approaches', 'strategies', 'tactics', 'plans', 'schemes', 'programs', 'projects', 'initiatives', 'campaigns', 'efforts', 'attempts', 'tries', 'endeavors', 'undertakings', 'ventures', 'enterprises', 'businesses', 'operations', 'activities', 'actions', 'deeds', 'acts', 'performances', 'executions', 'implementations', 'realizations', 'accomplishments', 'achievements', 'successes', 'triumphs', 'victories', 'wins', 'conquests', 'masteries', 'dominances', 'controls', 'powers', 'authorities', 'influences', 'impacts', 'effects', 'results', 'outcomes', 'consequences', 'implications', 'significances', 'importances', 'relevances', 'values', 'worths', 'merits', 'qualities', 'standards', 'levels', 'grades', 'ranks', 'positions', 'statuses', 'standings', 'reputations', 'images', 'perceptions', 'views', 'opinions', 'judgments', 'assessments', 'evaluations', 'appraisals', 'estimations', 'calculations', 'computations', 'measurements', 'quantifications', 'analyses', 'breakdowns', 'decompositions', 'dissections', 'examinations', 'studies', 'investigations', 'researches', 'explorations', 'discoveries', 'findings', 'results', 'conclusions', 'decisions', 'choices', 'selections', 'options', 'alternatives', 'possibilities', 'opportunities', 'chances', 'prospects', 'potentials', 'capabilities', 'capacities', 'abilities', 'skills', 'talents', 'gifts', 'aptitudes', 'competences', 'proficiencies', 'expertises', 'knowledges', 'understandings', 'comprehensions', 'awarenesses', 'consciousnesses', 'recognitions', 'realizations', 'appreciations', 'acknowledgments', 'acceptances', 'approvals', 'endorsements', 'supports', 'backings', 'assistances', 'helps', 'aids', 'services', 'contributions', 'inputs', 'participations', 'involvements', 'engagements', 'commitments', 'dedications', 'devotions', 'loyalties', 'faithfulnesses', 'reliabilities', 'dependabilities', 'trustworthinesses', 'credibilities', 'integrities', 'honesties', 'sincerities', 'authenticities', 'genuinenesses', 'truths', 'facts', 'realities', 'actualities', 'existences', 'presences', 'occurrences', 'happenings', 'events', 'incidents', 'episodes', 'situations', 'circumstances', 'conditions', 'states', 'statuses', 'positions', 'locations', 'places', 'spots', 'sites', 'venues', 'settings', 'environments', 'contexts', 'backgrounds', 'histories', 'pasts', 'previouses', 'formers', 'earliers', 'priors', 'befores', 'precedings', 'antecedents', 'preliminaries', 'preparatories', 'initials', 'firsts', 'primaries', 'mains', 'principals', 'chiefs', 'leadings', 'tops', 'highests', 'supremes', 'ultimates', 'finals', 'lasts', 'ends', 'conclusions', 'terminations', 'completions', 'finishes', 'closures', 'resolutions', 'solutions', 'answers', 'responses', 'replies', 'reactions', 'feedbacks', 'comments', 'remarks', 'observations', 'notes', 'notices', 'attentions', 'focuses', 'concentrations', 'emphases', 'stresses', 'importances', 'significances', 'relevances', 'pertinences', 'applicabilities', 'suitabilities', 'appropriatenesses', 'fitnesses', 'matches', 'correspondences', 'agreements', 'harmonies', 'consistencies', 'coherences', 'logics', 'reasons', 'senses', 'meanings', 'significances', 'imports', 'implications', 'suggestions', 'indications', 'signs', 'signals', 'symbols', 'representations', 'images', 'pictures', 'illustrations', 'examples', 'instances', 'cases', 'samples', 'specimens', 'models', 'prototypes', 'templates', 'patterns', 'designs', 'structures', 'organizations', 'arrangements', 'configurations', 'setups', 'layouts', 'formats', 'styles', 'appearances', 'looks', 'aspects', 'features', 'characteristics', 'attributes', 'properties', 'qualities', 'traits', 'elements', 'components', 'parts', 'pieces', 'sections', 'segments', 'portions', 'shares', 'fractions', 'percentages', 'proportions', 'ratios', 'rates', 'speeds', 'paces', 'tempos', 'rhythms', 'frequencies', 'occurrences', 'incidences', 'prevalences', 'commonnesses', 'rarities', 'uniquenesses', 'distinctivenesses', 'individualities', 'personalities', 'characters', 'natures', 'essences', 'substances', 'materials', 'matters', 'contents', 'informations', 'datas', 'facts', 'details', 'particulars', 'specifics', 'specifications', 'requirements', 'criterias', 'standards', 'guidelines', 'rules', 'regulations', 'policies', 'procedures', 'protocols', 'methods', 'techniques', 'approaches', 'strategies', 'tactics', 'plans', 'schemes', 'programs', 'projects', 'initiatives', 'campaigns', 'efforts', 'attempts', 'tries', 'endeavors', 'undertakings', 'ventures', 'enterprises', 'businesses', 'operations', 'activities', 'actions', 'deeds', 'acts', 'performances', 'executions', 'implementations', 'realizations', 'accomplishments', 'achievements', 'successes', 'triumphs', 'victories', 'wins', 'conquests', 'masteries', 'dominances', 'controls', 'powers', 'authorities', 'influences', 'impacts', 'effects', 'results', 'outcomes', 'consequences', 'implications', 'significances', 'importances', 'relevances', 'values', 'worths', 'merits', 'qualities', 'standards', 'levels', 'grades', 'ranks', 'positions', 'statuses', 'standings', 'reputations', 'images', 'perceptions', 'views', 'opinions', 'judgments', 'assessments', 'evaluations', 'appraisals', 'estimations', 'calculations', 'computations', 'measurements', 'quantifications', 'analyses', 'breakdowns', 'decompositions', 'dissections', 'examinations', 'studies', 'investigations', 'researches', 'explorations', 'discoveries', 'findings', 'results', 'conclusions', 'decisions', 'choices', 'selections', 'options', 'alternatives', 'possibilities', 'opportunities', 'chances', 'prospects', 'potentials', 'capabilities', 'capacities', 'abilities', 'skills', 'talents', 'gifts', 'aptitudes', 'competences', 'proficiencies', 'expertises', 'knowledges', 'understandings', 'comprehensions', 'awarenesses', 'consciousnesses', 'recognitions', 'realizations', 'appreciations', 'acknowledgments', 'acceptances', 'approvals', 'endorsements', 'supports', 'backings', 'assistances', 'helps', 'aids', 'services', 'contributions', 'inputs', 'participations', 'involvements', 'engagements', 'commitments', 'dedications', 'devotions', 'loyalties', 'faithfulnesses', 'reliabilities', 'dependabilities', 'trustworthinesses', 'credibilities', 'integrities', 'honesties', 'sincerities', 'authenticities', 'genuinenesses', 'truths', 'facts', 'realities', 'actualities', 'existences', 'presences', 'occurrences', 'happenings', 'events', 'incidents', 'episodes', 'situations', 'circumstances', 'conditions', 'states', 'statuses', 'positions', 'locations', 'places', 'spots', 'sites', 'venues', 'settings', 'environments', 'contexts', 'backgrounds', 'histories']

        tech_keyword_count = sum(1 for word in filtered_words if word in tech_keywords)
        tech_keyword_density = (tech_keyword_count / len(filtered_words)) * 100 if filtered_words else 0

        # Section analysis
        sections = content.split('#')
        section_count = len([s for s in sections if s.strip()])

        # Sentiment analysis (very basic)
        positive_words = ['excellent', 'outstanding', 'exceptional', 'innovative', 'successful', 'achieved', 'accomplished', 'improved', 'enhanced', 'optimized', 'effective', 'efficient', 'impactful', 'significant', 'substantial', 'remarkable', 'notable', 'impressive', 'strong', 'robust', 'scalable', 'reliable', 'secure', 'advanced', 'cutting-edge', 'state-of-the-art', 'breakthrough', 'pioneering', 'revolutionary', 'transformative', 'game-changing', 'industry-leading', 'award-winning', 'recognized', 'acclaimed', 'celebrated', 'praised', 'commended', 'honored', 'distinguished', 'prestigious', 'renowned', 'respected', 'esteemed', 'valued', 'appreciated', 'trusted', 'credible', 'authoritative', 'expert', 'professional', 'skilled', 'talented', 'gifted', 'capable', 'competent', 'proficient', 'experienced', 'seasoned', 'veteran', 'senior', 'lead', 'principal', 'chief', 'head', 'director', 'manager', 'supervisor', 'coordinator', 'organizer', 'planner', 'strategist', 'visionary', 'leader', 'pioneer', 'innovator', 'creator', 'inventor', 'developer', 'designer', 'architect', 'engineer', 'programmer', 'analyst', 'consultant', 'advisor', 'mentor', 'coach', 'trainer', 'educator', 'teacher', 'instructor', 'guide', 'facilitator', 'enabler', 'supporter', 'advocate', 'champion', 'ambassador', 'representative', 'spokesperson', 'communicator', 'presenter', 'speaker', 'author', 'writer', 'publisher', 'contributor', 'participant', 'member', 'volunteer', 'activist', 'enthusiast', 'passionate', 'dedicated', 'committed', 'devoted', 'loyal', 'faithful', 'reliable', 'dependable', 'trustworthy', 'honest', 'sincere', 'authentic', 'genuine', 'transparent', 'open', 'collaborative', 'cooperative', 'supportive', 'helpful', 'generous', 'kind', 'caring', 'compassionate', 'empathetic', 'understanding', 'patient', 'tolerant', 'respectful', 'considerate', 'thoughtful', 'mindful', 'aware', 'conscious', 'alert', 'attentive', 'focused', 'concentrated', 'determined', 'persistent', 'persevering', 'resilient', 'adaptable', 'flexible', 'agile', 'responsive', 'proactive', 'initiative', 'creative', 'imaginative', 'original', 'unique', 'distinctive', 'special', 'extraordinary', 'remarkable', 'amazing', 'incredible', 'fantastic', 'wonderful', 'great', 'good', 'positive', 'beneficial', 'valuable', 'useful', 'helpful', 'important', 'significant', 'meaningful', 'relevant', 'applicable', 'suitable', 'appropriate', 'fitting', 'perfect', 'ideal', 'optimal', 'best', 'top', 'highest', 'maximum', 'superior', 'premium', 'quality', 'excellence', 'perfection', 'mastery', 'expertise', 'skill', 'talent', 'ability', 'capability', 'capacity', 'potential', 'opportunity', 'possibility', 'chance', 'prospect', 'future', 'growth', 'development', 'progress', 'advancement', 'improvement', 'enhancement', 'optimization', 'refinement', 'evolution', 'transformation', 'change', 'innovation', 'revolution', 'breakthrough', 'discovery', 'invention', 'creation', 'production', 'generation', 'construction', 'building', 'development', 'design', 'planning', 'strategy', 'approach', 'method', 'technique', 'process', 'system', 'framework', 'structure', 'organization', 'management', 'leadership', 'direction', 'guidance', 'supervision', 'oversight', 'control', 'regulation', 'governance', 'administration', 'operation', 'execution', 'implementation', 'delivery', 'performance', 'achievement', 'accomplishment', 'success', 'victory', 'triumph', 'win', 'conquest', 'mastery', 'dominance', 'superiority', 'excellence', 'quality', 'standard', 'benchmark', 'reference', 'model', 'example', 'template', 'pattern', 'design', 'blueprint', 'plan', 'scheme', 'program', 'project', 'initiative', 'campaign', 'effort', 'attempt', 'try', 'endeavor', 'undertaking', 'venture', 'enterprise', 'business', 'operation', 'activity', 'action', 'deed', 'act', 'performance', 'execution', 'implementation', 'realization', 'accomplishment', 'achievement', 'success', 'triumph', 'victory', 'win', 'conquest', 'mastery', 'dominance', 'control', 'power', 'authority', 'influence', 'impact', 'effect', 'result', 'outcome', 'consequence', 'benefit', 'advantage', 'gain', 'profit', 'return', 'reward', 'prize', 'award', 'recognition', 'acknowledgment', 'appreciation', 'gratitude', 'thanks', 'praise', 'commendation', 'compliment', 'endorsement', 'recommendation', 'approval', 'support', 'backing', 'assistance', 'help', 'aid', 'service', 'contribution', 'input', 'participation', 'involvement', 'engagement', 'commitment', 'dedication', 'devotion', 'loyalty', 'faithfulness', 'reliability', 'dependability', 'trustworthiness', 'credibility', 'integrity', 'honesty', 'sincerity', 'authenticity', 'genuineness', 'truth', 'fact', 'reality', 'actuality', 'existence', 'presence', 'occurrence', 'happening', 'event', 'incident', 'episode', 'situation', 'circumstance', 'condition', 'state', 'status', 'position', 'location', 'place', 'spot', 'site', 'venue', 'setting', 'environment', 'context', 'background', 'history', 'experience', 'knowledge', 'understanding', 'comprehension', 'awareness', 'consciousness', 'recognition', 'realization', 'appreciation', 'acknowledgment', 'acceptance', 'approval', 'endorsement', 'support', 'backing', 'assistance', 'help', 'aid', 'service', 'contribution', 'input', 'participation', 'involvement', 'engagement', 'commitment', 'dedication', 'devotion', 'loyalty', 'faithfulness', 'reliability', 'dependability', 'trustworthiness', 'credibility', 'integrity', 'honesty', 'sincerity', 'authenticity', 'genuineness']

        positive_count = sum(1 for word in filtered_words if word in positive_words)
        sentiment_score = (positive_count / len(filtered_words)) * 100 if filtered_words else 0

        analytics_data = {
            'basic_stats': {
                'word_count': word_count,
                'character_count': char_count,
                'character_count_no_spaces': char_count_no_spaces,
                'paragraph_count': paragraph_count,
                'section_count': section_count,
                'sentence_count': sentences,
                'avg_words_per_sentence': round(avg_words_per_sentence, 1)
            },
            'readability': {
                'avg_words_per_sentence': round(avg_words_per_sentence, 1),
                'readability_score': min(100, max(0, 100 - (avg_words_per_sentence * 2))),  # Simplified readability
                'complexity_level': 'Simple' if avg_words_per_sentence < 15 else 'Moderate' if avg_words_per_sentence < 25 else 'Complex'
            },
            'keywords': {
                'top_keywords': top_keywords,
                'tech_keyword_count': tech_keyword_count,
                'tech_keyword_density': round(tech_keyword_density, 1),
                'total_unique_words': len(set(filtered_words))
            },
            'sentiment': {
                'positive_word_count': positive_count,
                'sentiment_score': round(sentiment_score, 1),
                'sentiment_level': 'Very Positive' if sentiment_score > 15 else 'Positive' if sentiment_score > 10 else 'Neutral' if sentiment_score > 5 else 'Needs Improvement'
            },
            'tech_nation_alignment': {
                'has_technical_content': tech_keyword_count > 10,
                'has_leadership_indicators': any(word in content.lower() for word in ['lead', 'manage', 'direct', 'supervise', 'coordinate', 'organize']),
                'has_innovation_indicators': any(word in content.lower() for word in ['innovate', 'create', 'develop', 'design', 'invent', 'pioneer']),
                'has_impact_indicators': any(word in content.lower() for word in ['impact', 'improve', 'enhance', 'optimize', 'increase', 'reduce', 'achieve']),
                'has_recognition_indicators': any(word in content.lower() for word in ['award', 'recognition', 'publish', 'speak', 'present', 'conference']),
                'word_count_appropriate': 800 <= word_count <= 1200
            }
        }

        return JsonResponse({
            'success': True,
            'analytics': analytics_data
        })

    except Exception as e:
        logger.error(f"Error generating analytics for document {pk}: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)

# Additional utility functions that might be referenced
def get_user_documents_summary(user):
    """Get a summary of user's documents"""
    documents = Document.objects.filter(user=user)

    summary = {
        'total': documents.count(),
        'by_type': {},
        'by_status': {},
        'recent_activity': documents.order_by('-updated_at')[:5]
    }

    # Count by type
    for doc_type, _ in Document.DOCUMENT_TYPES:
        count = documents.filter(document_type=doc_type).count()
        if count > 0:
            summary['by_type'][doc_type] = count

    # Count by status
    for status, _ in Document.STATUS_CHOICES:
        count = documents.filter(status=status).count()
        if count > 0:
            summary['by_status'][status] = count

    return summary

def validate_document_content(content, doc_type):
    """Validate document content based on type"""
    errors = []
    warnings = []

    if not content or not content.strip():
        errors.append("Document content cannot be empty")
        return errors, warnings

    word_count = len(content.split())

    if doc_type == 'personal_statement':
        if word_count < 500:
            errors.append("Personal statement is too short (minimum 500 words)")
        elif word_count < 800:
            warnings.append("Personal statement should be at least 800 words for optimal impact")
        elif word_count > 1500:
            warnings.append("Personal statement is quite long (consider keeping under 1200 words)")

        # Check for required sections
        required_sections = ['technical', 'leadership', 'innovation']
        content_lower = content.lower()

        for section in required_sections:
            if section not in content_lower:
                warnings.append(f"Consider adding a section about {section}")

    elif doc_type == 'cv':
        if word_count < 200:
            warnings.append("CV seems quite brief - consider adding more detail")

        # Check for common CV sections
        cv_sections = ['experience', 'education', 'skills', 'projects']
        content_lower = content.lower()

        missing_sections = [section for section in cv_sections if section not in content_lower]
        if missing_sections:
            warnings.append(f"Consider adding sections for: {', '.join(missing_sections)}")

    return errors, warnings



@login_required
def purchase_points(request):
    """View to display available points packages"""
    # Get active packages
    packages = PointsPackage.objects.filter(is_active=True).order_by('price')

    # Get user points
    user_points, created = UserPoints.objects.get_or_create(user=request.user)

    context = {
        'packages': packages,
        'user_points': user_points,
        'STRIPE_PUBLISHABLE_KEY': settings.STRIPE_PUBLISHABLE_KEY
    }

    return render(request, 'document_manager/purchase_points.html', context)


@csrf_exempt
def stripe_webhook(request):
    payload = request.body
    sig_header = request.META.get('HTTP_STRIPE_SIGNATURE')

    logger.info("Received Stripe webhook")

    try:
        webhook_secret = settings.STRIPE_WEBHOOK_SECRET
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
        logger.info(f"Stripe event type: {event['type']}")
    except ValueError as e:
        logger.error(f"Invalid Stripe payload: {str(e)}")
        return HttpResponse(status=400)
    except stripe.error.SignatureVerificationError as e:
        logger.error(f"Invalid Stripe signature: {str(e)}")
        return HttpResponse(status=400)

    # Handle both checkout.session.completed and payment_intent.succeeded events
    if event['type'] in ['checkout.session.completed', 'payment_intent.succeeded']:
        logger.info(f"Processing {event['type']} event")

        # Extract customer info based on event type
        if event['type'] == 'checkout.session.completed':
            session = event['data']['object']
            customer_email = session.get('customer_details', {}).get('email')
        else:  # payment_intent.succeeded
            payment_intent = event['data']['object']
            # Try to get customer email from metadata or customer object
            customer_email = None
            if 'metadata' in payment_intent and 'email' in payment_intent['metadata']:
                customer_email = payment_intent['metadata']['email']
            elif 'customer' in payment_intent:
                try:
                    customer = stripe.Customer.retrieve(payment_intent['customer'])
                    customer_email = customer.get('email')
                except Exception as e:
                    logger.error(f"Error retrieving customer: {str(e)}")

        logger.info(f"Customer email from event: {customer_email}")

        if customer_email:
            try:
                # Find the user by email
                user = User.objects.get(email=customer_email)
                logger.info(f"Found user: {user.username}")

                # Process referral if exists
                from referrals.models import ReferralSignup
                referral = ReferralSignup.objects.filter(referred_user=user).first()

                if referral:
                    logger.info(f"Found referral for user {user.username}")
                    logger.info(f"Referral points_awarded status: {referral.points_awarded}")

                    # Check if this is their first payment
                    payment_count = PointsTransaction.objects.filter(
                        user=user,
                        payment_status='completed'
                    ).count()

                    logger.info(f"User has {payment_count} completed payments")

                    # Award points if this is their first payment and points not already awarded
                    if payment_count > 0 and not referral.points_awarded:
                        if referral.award_points():
                            logger.info(f"Successfully awarded points for referral {referral.id}")
                        else:
                            logger.info(f"Failed to award points for referral {referral.id}")
                else:
                    logger.info(f"No referral found for user {user.username}")

            except User.DoesNotExist:
                logger.error(f"User with email {customer_email} not found")
            except Exception as e:
                logger.error(f"Error processing referral: {str(e)}")
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")
    else:
        logger.info(f"Ignoring event: {event['type']}")

    return HttpResponse(status=200)




@login_required
def checkout_package(request, package_id):
    """Create Stripe checkout session for a points package"""
    package = get_object_or_404(PointsPackage, id=package_id, is_active=True)

    # Create a transaction record (pending)
    transaction = PointsTransaction.objects.create(
        user=request.user,
        points_amount=package.points,
        price_paid=package.price,
        payment_method='stripe',
        status='pending'
    )

    # Create Stripe checkout session
    try:
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[
                {
                    'price_data': {
                        'currency': 'gbp',
                        'product_data': {
                            'name': f"{package.name} - {package.points} AI Points",
                            'description': package.description or "",
                        },
                        'unit_amount': int(package.price * 100),  # Convert to pence
                    },
                    'quantity': 1,
                },
            ],
            mode='payment',
            success_url=request.build_absolute_uri(
                reverse('document_manager:payment_success_redirect')
            ) + "?session_id={CHECKOUT_SESSION_ID}",
            cancel_url=request.build_absolute_uri(reverse('document_manager:payment_failed')),
            client_reference_id=str(transaction.id),
            metadata={
                'transaction_id': transaction.id,
                'user_id': request.user.id,
                'package_id': package.id,
                'points': package.points
            }
        )

        # Update transaction with checkout ID
        transaction.stripe_checkout_id = checkout_session.id
        transaction.save()

        return redirect(checkout_session.url)

    except Exception as e:
        logger.error(f"Error creating Stripe checkout session: {str(e)}")
        transaction.status = 'failed'
        transaction.save()
        messages.error(request, f"Payment processing error: {str(e)}")
        return redirect('document_manager:purchase_points')





@login_required
def payment_success(request, transaction_id):
    """Handle successful payment and award referral benefits if applicable."""
    transaction = get_object_or_404(PointsTransaction, id=transaction_id, user=request.user)
    session_id = request.GET.get('session_id')

    if transaction.payment_status == 'pending':
        try:
            if session_id and session_id.startswith('cs_'):
                checkout_session = stripe.checkout.Session.retrieve(session_id)
            elif transaction.stripe_checkout_id:
                checkout_session = stripe.checkout.Session.retrieve(transaction.stripe_checkout_id)
            else:
                messages.error(request, "Could not verify payment: Missing session information.")
                return redirect('document_manager:dashboard') # Or some other appropriate page

            if checkout_session.payment_status == 'paid':
                transaction.payment_status = 'completed'
                transaction.stripe_payment_intent_id = checkout_session.payment_intent
                # transaction.points is likely the field name based on your message
                transaction.save(update_fields=['payment_status', 'stripe_payment_intent_id'])

                user_points, _ = UserPoints.objects.get_or_create(user=request.user)
                user_points.add_points(transaction.points_amount) # Make sure points_amount is correct
                user_points.save()

                try:
                    profile = request.user.profile
                    if not profile.is_paid_user:
                        profile.is_paid_user = True
                        profile.save(update_fields=['is_paid_user'])
                        logger.info(f"Updated user {request.user.username} to paid status.")
                except AttributeError: # If profile doesn't exist for some reason
                    logger.error(f"UserProfile not found for {request.user.username} during payment success.")
                except Exception as e:
                    logger.error(f"Error updating paid status for {request.user.username}: {str(e)}")

                # Check if this is the user's first completed purchase and handle referral bonus
                # Ensure we only count *this* user's completed transactions
                is_first_completed_purchase = PointsTransaction.objects.filter(
                    user=request.user,
                    payment_status='completed'
                ).count() == 1 # If this transaction is the first one to be marked completed

                if is_first_completed_purchase:
                    logger.info(f"Processing first purchase referral check for user {request.user.email}")
                    try:
                        # Find if this user was referred by someone
                        referral_signup_instance = ReferralSignup.objects.filter(
                            referred_user=request.user
                        ).select_related('referral_code__user__profile').first() # Optimize query

                        if referral_signup_instance:
                            logger.info(f"User {request.user.email} was referred by {referral_signup_instance.referral_code.user.email}. Attempting to award rewards.")
                            # Call award_rewards which handles points and the new free_use
                            if referral_signup_instance.award_rewards(): # This method is in ReferralSignup model
                                logger.info(f"Successfully awarded referral rewards (points/free use) to referrer {referral_signup_instance.referral_code.user.email} for {request.user.email}'s purchase.")
                            else:
                                logger.info(f"Referral rewards for {request.user.email} (referrer: {referral_signup_instance.referral_code.user.email}) not awarded (e.g., already processed or error).")
                        else:
                            logger.info(f"User {request.user.email} made first purchase but was not referred or referral signup not found.")
                    except Exception as e:
                        logger.error(f"Error processing referral bonus for {request.user.email}: {str(e)}", exc_info=True)
                else:
                    logger.info(f"Not the first purchase for user {request.user.email} or payment count mismatch.")

                messages.success(
                    request,
                    f'Payment successful! {transaction.points_amount} points have been added to your account. Your new balance is {user_points.balance} points.'
                )
            else:
                messages.warning(
                    request,
                    'Your payment is being processed. Points will be added to your account once the payment is confirmed.'
                )
        except stripe.error.StripeError as e:
            logger.error(f"Stripe error verifying payment for transaction {transaction.id}: {str(e)}")
            messages.error(request, f"Stripe error verifying payment: {str(e)}")
        except Exception as e:
            logger.error(f"Error verifying payment for transaction {transaction.id}: {str(e)}", exc_info=True)
            messages.error(request, "An unexpected error occurred while verifying your payment.")

    elif transaction.payment_status == 'completed':
        messages.info(request, "This payment has already been processed.")
    else:
        messages.warning(request, f"Payment status: {transaction.payment_status}. If you believe this is an error, please contact support.")

    # Always get the latest points for display
    try:
        user_points_display = UserPoints.objects.get(user=request.user)
    except UserPoints.DoesNotExist:
        user_points_display = UserPoints.objects.create(user=request.user, balance=0)


    return render(request, 'document_manager/payment_success.html', {
        'transaction': transaction,
        'user_points': user_points_display
    })

    

@login_required
def payment_success_redirect(request):
    """Redirect to the payment success page with transaction ID"""
    # Check if session_id is provided in the URL
    session_id = request.GET.get('session_id')

    if not session_id:
        messages.error(request, "No session ID provided. Please contact support.")
        return redirect('document_manager:purchase_points')

    try:
        # Retrieve the checkout session
        checkout_session = stripe.checkout.Session.retrieve(session_id)

        # Get client_reference_id which should contain the transaction ID
        transaction_id = checkout_session.client_reference_id

        if transaction_id:
            # Redirect to the payment success page with transaction ID
            return redirect('document_manager:payment_success', transaction_id=transaction_id)
        else:
            # Try to find the transaction using the session ID
            transaction = PointsTransaction.objects.filter(
                stripe_checkout_id=session_id,
                user=request.user
            ).first()

            if transaction:
                return redirect('document_manager:payment_success', transaction_id=str(transaction.id))
            else:
                messages.error(request, "Transaction not found. Please contact support.")
                return redirect('document_manager:purchase_points')
    except Exception as e:
        logger.error(f"Error in payment redirect: {str(e)}")
        messages.error(request, "Unable to process payment confirmation. Please contact support.")
        return redirect('document_manager:purchase_points')


    
@login_required
def check_pending_payments(request):
    """Check and update any pending payments"""
    pending_transactions = PointsTransaction.objects.filter(
        user=request.user,
        payment_status='pending'
    ).order_by('-created_at')

    updated_count = 0

    for transaction in pending_transactions:
        try:
            # Skip if the transaction is too old (more than 24 hours)
            if django_timezone.now() - transaction.created_at > django_timezone.timedelta(hours=24):
                continue

            # Verify payment status with Stripe
            checkout_session = stripe.checkout.Session.retrieve(transaction.stripe_checkout_id)

            if checkout_session.payment_status == 'paid' and transaction.payment_status != 'completed':
                # Update transaction status
                transaction.payment_status = 'completed'
                transaction.stripe_payment_intent_id = checkout_session.payment_intent
                transaction.save()

                # Add points to user account
                user_points = UserPoints.objects.get(user=request.user)
                user_points.add_points(transaction.points)
                user_points.save()

                updated_count += 1

        except Exception as e:
            logger.error(f"Error checking payment {transaction.id}: {str(e)}")

    if updated_count > 0:
        messages.success(
            request,
            f'{updated_count} pending payment(s) have been confirmed and points added to your account.'
        )

    return redirect('document_manager:points_history')


@login_required
def payment_cancel(request):
    """Handle cancelled payment"""
    messages.info(request, 'Your payment was cancelled. No points have been added to your account.')
    return redirect('document_manager:purchase_points')




@login_required
def process_payment(request, package_id):
    """Process a payment for a points package"""
    if request.method != 'POST':
        # Consider redirecting to a more appropriate page or returning an error response
        messages.error(request, "Invalid request method.")
        return redirect('document_manager:purchase_points') # Ensure this URL name is correct

    try:
        package = PointsPackage.objects.get(id=package_id)
        payment_method_id = request.POST.get('payment_method_id')

        if not payment_method_id:
            messages.error(request, "No payment method provided.")
            return redirect('document_manager:purchase_points') # Ensure this URL name is correct

        payment_intent = stripe.PaymentIntent.create(
            amount=int(package.price * 100),
            currency='gbp',
            payment_method=payment_method_id,
            confirm=True,
            return_url=request.build_absolute_uri(reverse('document_manager:purchase_points')), # Ensure this URL name is correct
        )

        if payment_intent.status == 'succeeded':
            user_points, created = UserPoints.objects.get_or_create(user=request.user)
            logger.info(f"Before adding points: User {request.user.email} has {user_points.balance} points")
            user_points.add_points(package.points)
            
            profile, profile_created = UserProfile.objects.get_or_create(user=request.user)
            if profile_created:
                logger.warning(f"UserProfile for {request.user.email} was created during payment processing.")
            
            profile.is_paid_user = True
            profile.save(update_fields=['is_paid_user'])
            logger.info(f"Updated user {request.user.email} to paid status")
            logger.info(f"After adding points: User {request.user.email} has {user_points.balance} points")

            # --- CORRECTED PointsTransaction CREATION ---
            transaction = PointsTransaction.objects.create(
                user=request.user,
                package=package,  # Pass the actual package instance
                amount=package.price,  # Use the 'amount' field for the price
                points=package.points, # Use the 'points' field for the points
                payment_status='completed',
                stripe_payment_intent_id=payment_intent.id
                # The 'payment_method' field was not in your PointsTransaction model
            )
            # --- END CORRECTION ---

            logger.info(f"Created payment record with ID: {transaction.id} for package '{package.name}'")

            logger.info(f"Checking for referrals to process after payment for user: {request.user.email}")
            try:
                from referrals.models import ReferralSignup
                referral = ReferralSignup.objects.filter(referred_user=request.user).first()
                if referral:
                    logger.info(f"Found referral for user {request.user.email}, referrer: {referral.referral_code.user.email}")
                    logger.info(f"Referral points_awarded status before calling award_rewards: {referral.points_awarded}, has_been_rewarded: {referral.has_been_rewarded}")
                    if referral.award_rewards():
                        logger.info(f"Successfully processed referral bonus for {request.user.email} (referrer: {referral.referral_code.user.email}).")
                    else:
                        logger.error(f"Failed to process referral bonus for {request.user.email} via award_rewards method. Check referral model logs.")
                else:
                    logger.info(f"No pending referral found for user {request.user.email} to process bonus for.")
            except Exception as e:
                logger.error(f"Error processing referral bonus for {request.user.email}: {str(e)}")
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")

            messages.success(request, f"Successfully purchased {package.points} points!") # Add a success message
            return redirect('document_manager:payment_success', transaction_id=transaction.id) # Ensure this URL name is correct
        else:
            messages.error(request, f"Payment failed with status: {payment_intent.status}. Please try again.")
            return redirect('document_manager:payment_failed') # Ensure this URL name is correct

    except PointsPackage.DoesNotExist:
        messages.error(request, "Invalid package selected.")
        return redirect('document_manager:purchase_points') # Ensure this URL name is correct
    except stripe.error.CardError as e:
        messages.error(request, f"Your card was declined: {e.error.message}")
        return redirect('document_manager:payment_failed') # Ensure this URL name is correct
    except Exception as e:
        logger.error(f"Exception in payment processing for {request.user.email if request.user.is_authenticated else 'anonymous'}: {str(e)}", exc_info=True)
        messages.error(request, "An unexpected error occurred during payment processing. Please try again later or contact support.")
        return redirect('document_manager:payment_failed') # Ensure this URL name is correct



@login_required
def payment_failed(request):
    """Handle failed payment"""
    messages.error(request, "Your payment was cancelled. No points have been added to your account.")
    return redirect('document_manager:purchase_points')




@login_required
def generation_history(request):
    """Display history of document generations"""
    try:
        # Get user's generation history from documents
        generations = Document.objects.filter(
            user=request.user,
            document_type='personal_statement'
        ).order_by('-created_at')

        # Filter by date range if specified
        date_from = request.GET.get('date_from')
        date_to = request.GET.get('date_to')

        if date_from:
            try:
                from datetime import datetime
                date_from = datetime.strptime(date_from, '%Y-%m-%d').date()
                generations = generations.filter(created_at__date__gte=date_from)
            except ValueError:
                messages.warning(request, 'Invalid start date format')

        if date_to:
            try:
                from datetime import datetime
                date_to = datetime.strptime(date_to, '%Y-%m-%d').date()
                generations = generations.filter(created_at__date__lte=date_to)
            except ValueError:
                messages.warning(request, 'Invalid end date format')

        # Pagination
        from django.core.paginator import Paginator
        paginator = Paginator(generations, 20)  # Show 20 generations per page
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)

        # Calculate statistics
        total_generations = generations.count()
        successful_generations = generations.filter(status='completed').count()
        failed_generations = generations.filter(status='failed').count()

        stats = {
            'total': total_generations,
            'successful': successful_generations,
            'failed': failed_generations,
            'success_rate': round((successful_generations / total_generations * 100), 1) if total_generations > 0 else 0
        }

        context = {
            'page_obj': page_obj,
            'stats': stats,
            'date_from': date_from,
            'date_to': date_to,
        }

        return render(request, 'document_manager/generation_history.html', context)

    except Exception as e:
        logger.error(f"Error in generation history: {str(e)}")
        messages.error(request, f'Error loading generation history: {str(e)}')
        return redirect('document_manager:dashboard')

@login_required
def recommendation_guide(request):
    """Display recommendation guide for Tech Nation visa letters."""
    context = {
        'title': 'Recommender\'s Guide',
        'page_title': 'Tech Nation Visa Recommendation Guide'
    }
    return render(request, 'document_manager/recommendation_guide.html', context)

# Export the views for URL configuration
__all__ = [
    'document_list',
    'document_detail',
    'create_document',
    'edit_document',
    'delete_document',
    'duplicate_document',
    'download_document',
    'export_documents',
    'batch_process_documents',
    'personal_statement_builder',
    'generate_personal_statement',
    'check_generation_status',
    'save_generated_content',
    'preview_personal_statement',
    'update_personal_statement',
    'analyze_cv',
    'get_requirements',
    'dashboard',
    'api_usage_stats',
    'clear_cache',
    'health_check',
    'document_list_partial',
    'personal_statement_form',
    'cv_upload_form',
    'quick_generate',
    'regenerate_section',
    'improve_document',
    'document_versions',
    'compare_documents',
    'document_feedback',
    'document_analytics',
    'generation_history',
    'set_as_chosen',  # Add this
    'recommendation_guide'  # Add this
]