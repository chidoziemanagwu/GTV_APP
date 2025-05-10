# document_manager/services.py

from google import genai
from openai import OpenAI  # Add OpenAI
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
import os
import re
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class GeminiDocumentGenerator:
    def __init__(self):
        # Get API keys from environment variables
        gemini_api_key = os.getenv('GEMINI_API_KEY')
        openai_api_key = os.getenv('OPENAI_API_KEY')

        if not gemini_api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")
        if not openai_api_key:
            raise ValueError("OPENAI_API_KEY not found in environment variables")

        # Initialize both clients
        self.gemini_client = genai.Client(api_key=gemini_api_key)
        self.openai_client = OpenAI(api_key=openai_api_key)

        self.DISCLAIMER_TEXT = """IMPORTANT DISCLAIMER:

        This document is auto-generated using AI for SAMPLE PURPOSES ONLY.

        1. This is NOT an official Tech Nation document
        2. This is meant to serve as a starting point/template only
        3. You MUST heavily customize this content based on your personal experiences
        4. Do NOT submit this document as-is with your visa application
        5. Use this as a reference to understand the format and requirements
        6. Replace all content with your own authentic experiences and achievements
        7. Seek professional guidance for your final application

        The Tech Nation Visa Assistant platform provides these AI-generated templates
        to help you understand the expected format and content structure.
        The actual content must be personalized to reflect your unique journey and contributions.
        """

    def generate_personal_statement(self, cv_content, instructions=None):
        """Generate personal statement using OpenAI based on CV content"""
        system_prompt = """You are an expert in Tech Nation Global Talent Visa applications.
        Generate a compelling personal statement following this exact structure:

        1. Title: "Personal Statement for Tech Nation Global Talent Visa"

        2. Organize into these sections with clear headings:
        - Introduction
        - Technical Expertise and Innovation
        - Impact and Achievements
        - Contribution to UK Tech Sector
        - Leadership and Community
        - Future Vision
        - Conclusion

        Requirements:
        - Make section headings bold using "**Section Name:**"
        - Use bullet points (•) for listing achievements or skills
        - Any placeholder text that needs user input should be in [highlighted format]
        - Keep sections clearly separated with line breaks
        - Total length: 800-1000 words
        - Focus on technical excellence, measurable impact, and future potential
        - Use specific metrics and achievements where possible

        Important: Remove any meta-commentary or notes about the generation process.
        Start directly with the personal statement content."""

        user_prompt = f"""Generate a personal statement based on this CV and instructions:

        CV Content:
        {cv_content}

        Additional Instructions:
        {instructions if instructions else 'Focus on technical achievements and innovation'}"""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )

            if response.choices and response.choices[0].message.content:
                content = response.choices[0].message.content
                formatted_content = self._format_content(content)
                return formatted_content
            return None

        except Exception as e:
            print(f"Error generating personal statement: {str(e)}")
            # Fallback to Gemini if OpenAI fails
            return self._generate_with_gemini(cv_content, instructions)

    def _generate_with_gemini(self, cv_content, instructions=None):
        """Fallback method using Gemini"""
        try:
            prompt = f"""You are an expert in Tech Nation Global Talent Visa applications.
            [... your existing Gemini prompt ...]"""

            response = self.gemini_client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt
            )

            if response and hasattr(response, 'text'):
                return self._format_content(response.text)
            return None
        except Exception as e:
            print(f"Error with Gemini fallback: {str(e)}")
            return None

    def _format_content(self, content):
        """Format the content with HTML and CSS classes"""
        try:
            # Remove any meta-commentary before the actual content
            content = re.sub(r'^\*\*\s*$', '', content, flags=re.MULTILINE)  # Remove standalone **
            content = content.strip()

            # Format section headers
            content = re.sub(
                r'^([A-Za-z\s]+):$',
                r'<h3 class="text-xl font-bold text-gray-900 mt-8 mb-4">\1:</h3>',
                content,
                flags=re.MULTILINE
            )

            # Format bullet points and create proper lists
            lines = content.split('\n')
            formatted_lines = []
            in_list = False
            list_items = []

            for line in lines:
                line = line.strip()
                if not line:
                    if in_list and list_items:
                        formatted_lines.append(
                            f'<ul class="list-disc pl-6 space-y-2 mb-4">{"".join(list_items)}</ul>'
                        )
                        in_list = False
                        list_items = []
                    formatted_lines.append('<div class="my-4"></div>')
                    continue

                # Check if line starts with bullet point
                if line.startswith('•'):
                    if not in_list:
                        in_list = True
                        list_items = []
                    content_after_bullet = line[1:].strip()
                    if content_after_bullet:
                        list_items.append(
                            f'<li class="ml-4">{content_after_bullet}</li>'
                        )
                else:
                    if in_list and list_items:
                        formatted_lines.append(
                            f'<ul class="list-disc pl-6 space-y-2 mb-4">{"".join(list_items)}</ul>'
                        )
                        in_list = False
                        list_items = []

                    if not line.startswith('<h3'):
                        formatted_lines.append(f'<p class="text-gray-700 mb-4">{line}</p>')
                    else:
                        formatted_lines.append(line)

            # Handle any remaining list items
            if in_list and list_items:
                formatted_lines.append(
                    f'<ul class="list-disc pl-6 space-y-2 mb-4">{"".join(list_items)}</ul>'
                )

            # Combine into final HTML
            formatted_content = f"""
                <div class="personal-statement prose max-w-none bg-white rounded-lg shadow-sm p-8">
                    <div class="mb-8 border-b pb-6">
                        <h1 class="text-3xl font-bold text-gray-900 mb-4">Personal Statement Sample</h1>
                        <p class="text-sm text-gray-600">Generated for Tech Nation Global Talent Visa Application</p>
                        <p class="text-xs text-red-500 mt-2">This is a sample document - customize before use</p>
                    </div>
                    <div class="content space-y-4">
                        {''.join(formatted_lines)}
                    </div>
                    <div class="mt-8 pt-6 border-t">
                        <p class="text-sm text-gray-500 italic">Generated on: {datetime.now().strftime('%d %B %Y')}</p>
                        <p class="text-xs text-red-500 mt-2">Important: This is a template only - must be customized for actual visa application</p>
                    </div>
                </div>
            """

            return formatted_content

        except Exception as e:
            print(f"Error formatting content: {str(e)}")
            return content    
    
    
    def enhance_cv(self, original_cv, instructions=None):
        """Enhance CV using Gemini"""
        prompt = f"""You are an expert CV writer specializing in Tech Nation Global Talent Visa applications.
        Enhance the following CV to better align with Tech Nation visa criteria.
        Focus on:
        1. Technical achievements and innovation
        2. Leadership and impact
        3. Quantifiable results
        4. Contributions to the tech community
        5. Professional formatting

        Original CV:
        {original_cv}

        Additional Instructions:
        {instructions if instructions else 'Emphasize technical leadership and innovation'}"""

        try:
            response = self.client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt
            )
            return response.text if response and hasattr(response, 'text') else None
        except Exception as e:
            print(f"Error enhancing CV: {str(e)}")
            return None

    def generate_recommendation_guidelines(self, role, achievements):
        """Generate guidelines for recommendation letters"""
        prompt = f"""Create guidelines for a Tech Nation Global Talent Visa recommendation letter based on:
        Role: {role}
        Key Achievements: {achievements}

        Include:
        1. Suggested structure
        2. Key points to emphasize
        3. Technical achievements to highlight
        4. Impact metrics to mention"""

        try:
            response = self.client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt
            )
            return response.text if response and hasattr(response, 'text') else None
        except Exception as e:
            print(f"Error generating recommendation guidelines: {str(e)}")
            return None

    def save_to_docx(self, content, title, doc_type):
        """Save content to DOCX file with disclaimer"""
        try:
            doc = Document()

            # Add disclaimer with styling
            disclaimer = doc.add_paragraph()
            disclaimer_run = disclaimer.add_run("DISCLAIMER")
            disclaimer_run.bold = True
            disclaimer_run.font.color.rgb = RGBColor(255, 0, 0)

            # Add disclaimer text with proper formatting
            disclaimer_text = doc.add_paragraph(self.DISCLAIMER_TEXT)
            disclaimer_text.style = 'Intense Quote'

            # Add separator
            doc.add_paragraph("=" * 50)

            # Add title with proper styling
            title_heading = doc.add_heading(title, level=1)
            title_heading.style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue

            # Add generation date with styling
            date_paragraph = doc.add_paragraph()
            date_run = date_paragraph.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
            date_run.italic = True
            date_run.font.size = Pt(10)

            # Add separator before main content
            doc.add_paragraph("=" * 50)

            # Add main content with proper formatting
            content_paragraph = doc.add_paragraph(content)
            content_paragraph.style = 'Normal'

            # Create directory if it doesn't exist
            os.makedirs('generated_documents', exist_ok=True)

            # Generate filename with timestamp
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{doc_type}_{timestamp}.docx"
            filepath = os.path.join('generated_documents', filename)

            # Save the document
            doc.save(filepath)
            return filepath

        except Exception as e:
            print(f"Error saving to DOCX: {str(e)}")
            return None